"""
辩论备赛助手 — 一站式辩论备赛工具
"""
import streamlit as st
import os
import re
import sys
import json
import requests
from io import BytesIO
from bs4 import BeautifulSoup
from dotenv import load_dotenv, set_key, find_dotenv

# ============================================================
# 0. 初始化配置
# ============================================================
load_dotenv()
ENV_FILE = find_dotenv() or os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
SAVE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "saves")


def read_env(key, default=""):
    """读取配置：优先环境变量 → st.secrets（Streamlit Cloud）→ 默认值"""
    val = os.getenv(key)
    if val:
        return val
    try:
        import streamlit as st
        if hasattr(st, "secrets") and key in st.secrets:
            return st.secrets[key]
    except Exception:
        pass
    return default


def write_env(key, value, persist=True):
    """写入环境变量（内存中始终设置，文件写入失败时静默跳过）"""
    os.environ[key] = value
    if persist:
        try:
            if not os.path.exists(ENV_FILE):
                with open(ENV_FILE, "w") as f:
                    f.write("# 辩论备赛助手 配置文件\n")
            set_key(ENV_FILE, key, value)
        except Exception:
            pass  # 只读文件系统（如Streamlit Cloud）静默跳过


# ============================================================
# 1. LLM 客户端
# ============================================================
class LLMClient:
    def __init__(self, config: dict | None = None):
        cfg = config or {}
        self.provider = cfg.get("provider") or read_env("LLM_PROVIDER", "anthropic")
        self.api_key = cfg.get("api_key") or read_env("LLM_API_KEY") or read_env("ANTHROPIC_API_KEY")
        self.base_url = cfg.get("base_url") or read_env("LLM_BASE_URL", "https://api.anthropic.com/v1")
        self.model = cfg.get("model") or read_env("LLM_MODEL", "claude-sonnet-4-6")

        if self.provider == "anthropic":
            import anthropic
            self.client = anthropic.Anthropic(api_key=self.api_key)
        else:
            from openai import OpenAI
            self.client = OpenAI(api_key=self.api_key, base_url=self.base_url)

    def chat(self, messages: list[dict], temperature: float = 0.7) -> str:
        try:
            if self.provider == "anthropic":
                return self._chat_anthropic(messages, temperature)
            else:
                return self._chat_openai(messages, temperature)
        except Exception as e:
            return f"[LLM 调用错误] {e}"

    def _chat_anthropic(self, messages: list[dict], temperature: float) -> str:
        system = ""
        user_msgs = []
        for m in messages:
            if m["role"] == "system":
                system = m["content"]
            else:
                user_msgs.append({"role": m["role"], "content": m["content"]})
        resp = self.client.messages.create(
            model=self.model, max_tokens=8192, temperature=temperature,
            system=system or None, messages=user_msgs,
            thinking={"type": "disabled"},
        )
        parts = [b.text for b in resp.content if hasattr(b, "text")]
        return "\n".join(parts) or "[空响应]"

    def _chat_openai(self, messages: list[dict], temperature: float) -> str:
        resp = self.client.chat.completions.create(
            model=self.model, messages=messages,
            temperature=temperature, max_tokens=8192,
        )
        return resp.choices[0].message.content or ""

    def ask(self, system: str, user: str, temperature: float = 0.7) -> str:
        return self.chat([
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ], temperature)


_llm = None


def get_llm(config: dict | None = None) -> LLMClient:
    global _llm
    if config is not None:
        _llm = LLMClient(config)
    if _llm is None:
        _llm = LLMClient()
    return _llm


def reset_llm():
    global _llm
    _llm = None


# ============================================================
# 2. 网页搜索
# ============================================================
def has_exa() -> bool:
    return bool(read_env("EXA_API_KEY"))


def _exa_search(query: str, n: int = 5) -> list[dict]:
    try:
        resp = requests.post(
            "https://api.exa.ai/search",
            headers={"x-api-key": read_env("EXA_API_KEY"), "Content-Type": "application/json"},
            json={"query": query, "numResults": n, "useAutoprompt": True,
                  "contents": {"text": {"maxCharacters": 3000}}},
            timeout=30,
        )
        data = resp.json()
        return [{"title": r.get("title", ""), "url": r.get("url", ""), "text": r.get("text", "")}
                for r in data.get("results", [])]
    except Exception:
        return []


def fetch_page(url: str, max_chars: int = 5000) -> str:
    try:
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        resp = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 20]
        return "\n".join(lines)[:max_chars]
    except Exception:
        return ""


# ============================================================
# 3. 视频字幕
# ============================================================
def extract_video_captions(url: str) -> str:
    if "youtube.com" in url or "youtu.be" in url:
        return _youtube_captions(url)
    elif "bilibili.com" in url or "b23.tv" in url:
        return _bilibili_captions(url)
    return "[不支持的视频平台]"


def _youtube_captions(url: str) -> str:
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        m = re.search(r"(?:v=|/v/|youtu\.be/|embed/)([a-zA-Z0-9_-]{11})", url)
        if not m:
            return "[无法解析 YouTube ID]"
        tx = YouTubeTranscriptApi.get_transcript(m.group(1), languages=["zh", "en"])
        return " ".join([e["text"] for e in tx])
    except Exception as e:
        return f"[YouTube 字幕提取失败] {e}"


def _bilibili_captions(url: str) -> str:
    try:
        m = re.search(r"BV[a-zA-Z0-9]{10}", url)
        if not m:
            return "[无法解析 B站 BV 号]"
        bvid = m.group(0)
        headers = {"User-Agent": "Mozilla/5.0", "Referer": "https://www.bilibili.com"}
        info = requests.get(f"https://api.bilibili.com/x/web-interface/view?bvid={bvid}", headers=headers, timeout=15).json()
        if info.get("code") != 0:
            return f"[B站错误] {info.get('message','')}"
        data = info["data"]
        result = f"标题：{data.get('title','')}\n简介：{data.get('desc','')}\n"
        cid = data.get("cid", 0)
        if cid:
            sub_resp = requests.get(f"https://api.bilibili.com/x/player/v2?bvid={bvid}&cid={cid}", headers=headers, timeout=15).json()
            subs = sub_resp.get("data", {}).get("subtitle", {}).get("subtitles", [])
            if subs:
                sub_url = subs[0].get("subtitle_url", "")
                if sub_url:
                    if sub_url.startswith("//"):
                        sub_url = "https:" + sub_url
                    sub_json = requests.get(sub_url, headers=headers, timeout=15).json()
                    captions = " ".join([item.get("content", "") for item in sub_json.get("body", [])])
                    result += f"字幕：{captions}"
        return result
    except Exception as e:
        return f"[B站字幕提取失败] {e}"


# ============================================================
# 4. 文件读取
# ============================================================
def read_uploaded(uploaded) -> str:
    """读取 Streamlit UploadedFile 对象"""
    import tempfile
    suffix = os.path.splitext(uploaded.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded.read())
        path = tmp.name
    content = _read_file(path)
    try:
        os.unlink(path)
    except Exception:
        pass
    return content


def _read_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
            return "\n".join([p.extract_text() or "" for p in PdfReader(file_path).pages]).strip() or "[PDF 内容为空]"
        except Exception as e:
            return f"[PDF 读取错误] {e}"
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            return "\n".join([p.text for p in Document(file_path).paragraphs if p.text.strip()]).strip() or "[Word 内容为空]"
        except Exception as e:
            return f"[Word 读取错误] {e}"
    else:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk") as f:
                return f.read().strip()
        except Exception as e:
            return f"[文件读取错误] {e}"


# ============================================================
# 5. 资料研究员
# ============================================================
SYS_RESEARCHER = """你是一名资深辩论赛资料研究员。为辩手提炼最关键信息：
1. 概括各方核心论点和论证逻辑
2. 提取关键数据和事实，标注来源
3. 指出论证中的亮点和漏洞

请用结构化 Markdown 输出。"""


def do_research(topic: str, side: str, video_urls: list[str] | None = None) -> str:
    llm = get_llm()
    parts = []

    if has_exa():
        parts.append("## 一、过往辩论赛资料\n")
        for r in _exa_search(f"{topic} 辩论赛 文字记录 辩词", 5):
            parts.append(f"- **{r['title']}**\n  {r['text'][:600]}")
            extra = fetch_page(r["url"])
            if extra:
                parts.append(f"  {extra[:2000]}")
        parts.append("\n## 二、哲学与学术讨论\n")
        for r in _exa_search(f"{topic} 哲学 伦理学 论证分析", 4):
            parts.append(f"- **{r['title']}**\n  {r['text'][:600]}")
        parts.append("\n## 三、相关热点\n")
        for r in _exa_search(f"{topic} 最新新闻 社会热点", 4):
            parts.append(f"- **{r['title']}**\n  {r['text'][:600]}")
    else:
        parts.append(llm.ask(SYS_RESEARCHER,
            f"请针对辩题「{topic}」（持方：{side}），利用你的知识库生成结构化研究资料，"
            f"包括：过往辩论赛与论证策略、哲学与学术讨论、当下热点与案例。尽可能详细。",
            temperature=0.8))

    if video_urls:
        parts.append("\n## 四、视频资料字幕\n")
        for url in video_urls:
            parts.append(f"\n### {url}\n{extract_video_captions(url)[:5000]}")

    raw = "\n".join(parts)
    summary = llm.ask(SYS_RESEARCHER,
        f"辩题：{topic}\n持方：{side}\n\n以下是资料，请为{side}方提炼最关键的论证素材：\n\n{raw[:15000]}")
    return f"{summary}\n\n---\n\n### 原始资料附录\n{raw}"


# ============================================================
# 6. 备赛方案生成器
# ============================================================
SYS_PLAN = """你是传奇级辩论教练（曾带队获得国际大专辩论赛冠军）。基于研究资料制定备赛方案。

方案结构：
### 一、辩题解析
- 核心概念定义 / 关键争议点 / 双方论证义务

### 二、核心论点树（{side}方）
- 主论点 + 分论点 + 支撑数据/例子（3组）

### 三、对方攻击及反驳预案
- 对方可能的3个核心论点及各自反驳策略

### 四、关键数据卡
- 可引用的研究/统计 + 经典案例

### 五、哲学论证框架
- 适用的伦理学/哲学理论及运用方式

### 六、价值升华路径
- 从事实到价值的上升路径 + 结尾金句

方案必须具体可操作，直接给出能用于比赛的辩词片段。"""


def generate_plan(topic: str, side: str, research: str) -> str:
    llm = get_llm()
    return llm.ask(
        SYS_PLAN.replace("{side}", side),
        f"辩题：{topic}\n我方持方：{side}\n\n研究资料：\n{research[:12000]}\n\n请为{side}方制定完整备赛方案。",
        temperature=0.8,
    )


# ============================================================
# 7. 论点生成引擎
# ============================================================
SYS_ARGS = """你是辩论战略顾问，从哲学和现实两个维度构建论点。

每个论点：核心主张 + 哲学依据 + 现实例证 + 逻辑推理链 + ★评级
同时标注反方可能的攻击点和加固建议。"""


def generate_arguments(topic: str, side: str, research: str) -> str:
    llm = get_llm()
    hot_text = ""
    if has_exa():
        for r in _exa_search(f"{topic} 最新 2025 社会 案例 研究", 3):
            hot_text += f"- {r['title']}: {r['text'][:600]}\n"

    return llm.ask(SYS_ARGS,
        f"辩题：{topic}\n持方：{side}\n\n哲学框架参考：功利主义、义务论、权利论、契约论、美德伦理、关怀伦理\n\n"
        f"当下热点：{hot_text}\n\n既有研究：{research[:3000]}\n\n"
        f"请为{side}方生成5-7个核心论点，每个含哲学依据和现实例证。",
        temperature=0.85,
    )


# ============================================================
# 7.5 积木式立论块生成器
# ============================================================
SYS_BLOCKS = """你是一名辩论立论架构师。你的任务是为辩手生成立论搭建用的"积木块"。

生成JSON格式，包含五类积木块，每类3-5个选项：

{
  "approaches": [
    {"text": "开题思路：直接论证——直接与对方讨论辩题本身，从常识和直觉出发逐层深入"},
    ...
  ],
  "criteria": [
    {"text": "判准：应以最大多数人利益为根本判断标准（功利主义框架）"},
    ...
  ],
  "definitions": [
    {"text": "定义：「人工智能决策」指完全由AI系统自主做出的、无需人类干预的重大决策"},
    ...
  ],
  "arguments": [
    {"text": "论点：AI决策能消除人类的主观偏见，实现更公平的结果"},
    ...
  ],
  "logic_chains": [
    {"text": "逻辑链：效率优势→数据驱动优于直觉→减少医疗误诊→AI辅助诊断更可靠"},
    ...
  ]
}

生成要求：
- 开题思路（approaches）：这是立论的"破题"方式，提供3-5种不同的论证切入角度。包括但不限于：
  - 重新定义辩题中的核心概念（如"必需品"→"奢侈品"、"AI"→"工具延伸"）
  - 转换论证层次（从事实层转向价值层、从个体转向社会）
  - 设置论证前提（如"人没有必需品，所以爱情不是必需品"）
  - 类比对冲（用类似的、已被接受的命题来类比）
  - 划定边界法（明确哪些情况不在讨论范围）
  每个开题思路 40-150 字，给出具体的破题方向而非空泛策略。
- 判准（criteria）：不同的评价标准，决定从什么角度判断正反双方胜负
- 定义（definitions）：对辩题中关键概念的不同解读方式，每种定义引导不同的论证路径
- 论点（arguments）：完整的论点，含主张+理由+证据，直接可用于辩论
- 逻辑链（logic_chains）：从前提到结论的完整推理链条，用箭头连接各环节

每个积木块 60-200 字，内容精炼、独立可用。直接输出JSON，不要其他文字。"""


def generate_blocks(topic: str, side: str, research: str):
    """为辩题生成四类立论积木块"""
    llm = get_llm()
    result = llm.ask(
        SYS_BLOCKS,
        f"辩题：{topic}\n我方持方：{side}\n\n研究资料：\n{research[:5000]}\n\n请为{side}方生成四类立论积木块。直接输出JSON。",
        temperature=0.85,
    )
    # 解析 JSON
    try:
        # 提取 JSON（LLM 可能在前后加说明文字）
        json_start = result.find("{")
        json_end = result.rfind("}") + 1
        if json_start >= 0 and json_end > json_start:
            data = json.loads(result[json_start:json_end])
        else:
            raise ValueError("No JSON found")
    except Exception:
        # 解析失败时生成默认块
        data = {
            "criteria": [{"text": f"判准解析失败，请重试。原始输出：{result[:200]}"}],
            "definitions": [],
            "arguments": [],
            "logic_chains": [],
        }

    st.session_state.approach_blocks = [{"text": b["text"], "selected": False} for b in data.get("approaches", [])]
    st.session_state.criteria_blocks = [{"text": b["text"], "selected": False} for b in data.get("criteria", [])]
    st.session_state.definition_blocks = [{"text": b["text"], "selected": False} for b in data.get("definitions", [])]
    st.session_state.argument_blocks = [{"text": b["text"], "selected": False} for b in data.get("arguments", [])]
    st.session_state.logic_blocks = [{"text": b["text"], "selected": False} for b in data.get("logic_chains", [])]
    st.session_state.custom_blocks = []
    st.session_state.blocks_generated = True


# ============================================================
# 8. 分辩位训练器
# ============================================================
POSITION_PROFILES = {
    "一辩": {
        "role": "立论陈词手 · 前端接质",
        "duty": "一辩陈词：定义核心概念，搭建论证框架，展开主论点体系；前端接质询：守护我方框架根基",
        "style": "条理清晰、逻辑严密、定义精准、框架稳固",
        "system_prompt": """你是一辩教练。一辩承担两项核心任务：

【一辩陈词】（新国辩赛制3分钟）
1. 开场破题与核心概念定义（40秒）
2. 判准提出与论证框架搭建（40秒）
3. 展开2-3个主论点，每个含完整逻辑链（1.5分钟）
4. 小结，为后续环节预留接口（20秒）

要求：语言精准、概念无歧义；每个论点形成"主张-理由-证据"三段论；预判对方的定义争夺，提前锁定有利定义；开场有吸引力，结尾有力度。

【前端接质询】
一辩是对方三辩/四辩的首要质询目标。对方会攻击你的框架根基（定义、判准、前提）。
你必须像框架守护者一样守住核心概念，不能松动我方定义和判准。
接质要点：温和而坚定，正面回应，不反问。""",
    },
    "二辩": {
        "role": "申论驳论手 · 中端接质",
        "duty": "陈词：回应一辩遗留问题+申论深化+驳论攻击；中端接质询：应对战场核心交锋",
        "style": "反应敏捷、攻守兼备、逻辑拆解、灵活化解",
        "system_prompt": """你是二辩教练。二辩承担两项核心任务：

【二辩陈词】（新国辩赛制2分钟）
1. 回应对方一辩立论的核心论点，快速驳斥（30秒）
2. 申论深化：补充我方一辩未能展开的论证，引入新论据（1分钟）
3. 驳论攻击：逐条拆解对方论证的逻辑谬误（30秒）

要求：找逻辑谬误（偷换概念、以偏概全、因果倒置等）；用事实数据回击；强化一辩被攻击的薄弱点。

【中端接质询】
二辩申论驳论后被对方三辩质询，是战场核心交锋区。对方主要攻击你的论证深度和例证可靠性。
面对情景陷阱要灵活化解，善于用例子反制例子。接质不能反问，必须正面回应。""",
    },
    "三辩": {
        "role": "质询盘问手 · 中场收束",
        "duty": "质询对方中端核心战场；中场战场总结与论点概括收束",
        "style": "问题精准、连环追问、逻辑拆解、一剑封喉",
        "system_prompt": """你是三辩教练。三辩承担两项核心任务：

【质询对方中端核心战场】（新国辩赛制2.5分钟盘问+1.5分钟小结）
盘问策略：
- 设计3-5组问题链，每组3-4个递进问题
- 问题链终点指向对方论证在中端战场的逻辑矛盾
- 以封闭式问题为主（是/否），控制对方回答空间
- 策略：归谬法、类比法、特例法、二难推理

质询目标：对方的申论深化部分、对方的例证可靠性、对方的逻辑推理链

【中场战场总结与论点概括收束】（1.5分钟小结）
1. 梳理盘问中暴露的对方逻辑矛盾（30秒）
2. 归纳中端战场的核心分歧点（30秒）
3. 将分散的论点收束为我方论证体系的有机组成部分（30秒）

要求：问题简洁（不超过20字）；预判回答及追问方向；小结要有概括力和说服力。""",
    },
    "四辩": {
        "role": "全局质询手 · 价值收束",
        "duty": "质询对方前端问题（定义/判准/前提）；全局战场总结与价值升华收束",
        "style": "大局观强、价值引领、概括提炼、定音之锤",
        "system_prompt": """你是四辩教练。四辩承担两项核心任务：

【质询对方前端问题】（新国辩赛制：先质询后陈词）
质询目标：
- 攻击对方一辩立论中的定义是否自洽？
- 挑战对方判准是否合理、能否覆盖辩题的全部情况？
- 挖掘对方立论前提中的反例和漏洞
- 逼对方在核心概念上做出不利退让
问题风格：总结性强，帮评委看清对方框架的根本缺陷。

【全局战场总结与价值收束】（新国辩赛制3.5分钟）
1. 回顾全场争议焦点（1分钟）
   - 对方今天到底说了什么？核心逻辑链在哪里断裂？
   - 从一辩到三辩，对方在各个战场的表现如何？
2. 全局战场总结（1.5分钟）
   - 梳理前端（定义/判准）、中端（论证/攻防）两个战场的胜负
   - 将三辩的中场收束提升为全局结论
3. 价值升华（1分钟）
   - 从事实层面上升到价值层面
   - 告诉评委"为什么这个辩题重要"
   - 用金句收尾，留下最后也是最深的印象

要求：不做新论证，只梳理和升华；找对方最薄弱环节集中火力；语言有感染力和画面感；结尾要有"定音之锤"效果。""",
    },
}


class PositionTrainer:
    def __init__(self):
        self.llm = get_llm()

    def get_profile(self, pos: str) -> dict:
        return POSITION_PROFILES.get(pos, POSITION_PROFILES["一辩"])

    def train_from_material(self, pos: str, topic: str, side: str,
                            plan: str, material: str) -> str:
        p = self.get_profile(pos)
        return self.llm.ask(p["system_prompt"],
            f"辩题：{topic}\n我方持方：{side}\n辩位：{pos}（{p['role']}）\n\n"
            f"备赛方案：\n{plan[:4000]}\n\n上传资料：\n{material[:8000]}\n\n"
            f"请生成{pos}的完整辩词，严格按{pos}辩词结构输出，直接给出可背诵使用的辩词文本。",
            temperature=0.8,
        )

    def refine_input(self, pos: str, topic: str, side: str,
                     plan: str, draft: str, history: list[dict] | None = None) -> str:
        p = self.get_profile(pos)
        msgs = [{"role": "system", "content": p["system_prompt"]}]
        ctx = f"辩题：{topic}\n持方：{side}\n辩位：{pos}（{p['role']}）\n\n备赛方案：\n{plan[:3000]}"
        msgs.append({"role": "user", "content": f"【比赛背景】\n{ctx}"})
        msgs.append({"role": "assistant", "content": "好的，我已了解比赛背景。请发送你的想法或草稿，我来优化。"})
        if history:
            msgs.extend(history)
        msgs.append({"role": "user",
            "content": f"请将以下想法/草稿优化为{pos}正式辩词。保持核心观点，按{pos}风格完善语言、补充论证、优化结构：\n\n{draft}"})
        return self.llm.chat(msgs, temperature=0.8)


# ============================================================
# 9. Word 导出
# ============================================================
def generate_docx(title: str, content: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.styles["Normal"].font.size = Pt(12)
        doc.add_heading(f"{st.session_state.topic} — {st.session_state.side}方 {title}", level=0)
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                doc.add_heading(line.lstrip("# "), level=2)
            else:
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        st.error(f"导出失败：{e}")
        return b""


# ============================================================
# 10. Streamlit UI
# ============================================================
def init_session():
    defaults = {
        "page": "input", "topic": "", "side": "正方",
        "research_result": "", "plan": "", "arguments": "",
        "current_position": None, "chat_history": [],
        "position_results": {}, "video_url_str": "",
        "standpoint_confirmed": False, "user_arguments": [],
        # 积木式立论
        "blocks_generated": False,
        "approach_blocks": [],       # [{text, selected}]  开题思路
        "criteria_blocks": [],       # [{text, selected}]
        "definition_blocks": [],     # [{text, selected}]
        "argument_blocks": [],       # [{text, selected}]
        "logic_blocks": [],          # [{text, selected}]
        "custom_blocks": [],         # [{text, category}]
        # 框架审查与辩位任务
        "framework_review": "",      # 框架审查结果
        "framework_reviewed": False,
        "position_tasks": {},        # {一辩: "任务", 二辩: "...", ...}
        "position_speeches": {},     # {一辩: "初稿辩词", 二辩: "...", ...}
        # 接质训练
        "cross_exam_rounds": [],     # [{question, user_answer, ai_answer, followup, ...}]
        "splash_done": False,         # 启动画面是否已结束
        "cross_exam_current_q": "",  # 当前接质问题
        "cross_exam_need_help": False,
        "cross_exam_evaluation": "",  # 上一轮回答的评价
        "cross_exam_last_answer": "", # 上一轮的用户回答
        "editing_speech": False,      # 是否正在编辑辩词
        "speech_feedback": "",        # 编辑后的框架更新反馈
        # 三辩/四辩 质询训练（主动提问方）
        "interrogate_q": "",          # 用户当前要问的问题
        "interrogate_ai_a": "",       # AI扮演对方的回应
        "interrogate_eval": "",       # 问题质量评价
        "interrogate_history": [],    # 质询训练历史
        # 论据库
        "evidence_library": [],       # [{content, source, tags, added_at}]
        "evidence_notification": "",  # 论据库更新通知
        "_last_converted_speech": "", # 最近转化的辩词
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_session()

# ============================================================
# 0.5 会话保存/加载
# ============================================================
def _save_path(topic: str) -> str:
    """根据辩题生成安全的文件名"""
    safe = re.sub(r'[\\/*?:"<>|]', "_", topic)[:50]
    return os.path.join(SAVE_DIR, f"{safe}.json")


def save_session():
    """保存当前会话到文件"""
    if not st.session_state.topic:
        return
    os.makedirs(SAVE_DIR, exist_ok=True)
    data = {
        "topic": st.session_state.topic,
        "side": st.session_state.side,
        "research_result": st.session_state.research_result,
        "plan": st.session_state.plan,
        "arguments": st.session_state.arguments,
        "current_position": st.session_state.current_position,
        "position_results": st.session_state.position_results,
        "video_url_str": st.session_state.video_url_str,
        "standpoint_confirmed": st.session_state.standpoint_confirmed,
        "user_arguments": st.session_state.user_arguments,
        "chat_history": st.session_state.chat_history,
        "page": st.session_state.page,
        "approach_blocks": st.session_state.approach_blocks,
        "criteria_blocks": st.session_state.criteria_blocks,
        "definition_blocks": st.session_state.definition_blocks,
        "argument_blocks": st.session_state.argument_blocks,
        "logic_blocks": st.session_state.logic_blocks,
        "custom_blocks": st.session_state.custom_blocks,
        "blocks_generated": st.session_state.blocks_generated,
        "framework_review": st.session_state.framework_review,
        "framework_reviewed": st.session_state.framework_reviewed,
        "position_tasks": st.session_state.position_tasks,
        "position_speeches": st.session_state.position_speeches,
        "cross_exam_rounds": st.session_state.cross_exam_rounds,
        "interrogate_history": st.session_state.interrogate_history,
        "evidence_library": st.session_state.evidence_library,
    }
    with open(_save_path(st.session_state.topic), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_session(topic: str):
    """从文件恢复会话"""
    path = _save_path(topic)
    if not os.path.exists(path):
        return False
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    for k, v in data.items():
        st.session_state[k] = v
    return True


def list_saved_sessions() -> list[dict]:
    """列出所有已保存的会话"""
    if not os.path.exists(SAVE_DIR):
        return []
    sessions = []
    for fn in sorted(os.listdir(SAVE_DIR), reverse=True):
        if fn.endswith(".json"):
            path = os.path.join(SAVE_DIR, fn)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    d = json.load(f)
                mtime = os.path.getmtime(path)
                import datetime
                dt = datetime.datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M")
                sessions.append({
                    "topic": d.get("topic", fn.replace(".json", "")),
                    "side": d.get("side", ""),
                    "has_plan": bool(d.get("plan")),
                    "has_standpoint": bool(d.get("standpoint_confirmed")),
                    "page": d.get("page", "input"),
                    "saved_at": dt,
                })
            except Exception:
                pass
    return sessions


def delete_session(topic: str):
    """删除保存的会话"""
    path = _save_path(topic)
    if os.path.exists(path):
        os.remove(path)


st.markdown('''
<style>
    :root {
        --violet: #6C5CE7; --violet-ghost: #A29BFE; --violet-pale: #C4B5FD;
        --gold: #C9A96E; --text: #E8E6F0; --text-dim: #8A8A9A;
        --radius: 10px;
        --spring: cubic-bezier(0.34, 1.56, 0.64, 1);
        --smooth: cubic-bezier(0.25, 0.1, 0.25, 1);
    }
    section[data-testid="stAppViewContainer"] {
        background: radial-gradient(ellipse at 50% 0%, #14142A 0%, #0A0A0F 60%) !important;
    }
    header[data-testid="stHeader"] { background: transparent !important; }
    section[data-testid="stSidebar"] > div {
        background: rgba(10,10,15,0.88) !important;
        backdrop-filter: blur(40px) saturate(140%) !important;
        -webkit-backdrop-filter: blur(40px) saturate(140%) !important;
    }
    .stButton button, button[kind] {
        background: rgba(255,255,255,0.04) !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        border-radius: var(--radius) !important;
        color: var(--text) !important;
        font-size: 0.9rem !important; font-weight: 500 !important;
        transition: all 0.25s var(--smooth) !important;
        backdrop-filter: blur(10px) !important;
    }
    .stButton button:hover, button[kind]:hover {
        background: rgba(108,92,231,0.12) !important;
        border-color: rgba(108,92,231,0.35) !important;
        box-shadow: 0 4px 24px rgba(108,92,231,0.15) !important;
        transform: translateY(-1px) scale(1.01);
    }
    .stButton button:active, button[kind]:active {
        transform: scale(0.97) !important; transition: all 0.1s var(--smooth) !important;
    }
    button[kind="primary"] {
        background: linear-gradient(135deg, rgba(108,92,231,0.22), rgba(162,155,254,0.16)) !important;
        border-color: rgba(108,92,231,0.4) !important; font-weight: 600 !important;
    }
    button[kind="primary"]:hover {
        background: linear-gradient(135deg, rgba(108,92,231,0.35), rgba(162,155,254,0.25)) !important;
        box-shadow: 0 4px 32px rgba(108,92,231,0.25) !important;
    }
    input, textarea, [data-baseweb="input"] input, [data-baseweb="textarea"] textarea {
        background: rgba(255,255,255,0.03) !important;
        border: 1px solid rgba(255,255,255,0.08) !important;
        border-radius: var(--radius) !important;
        color: var(--text) !important;
        transition: all 0.35s var(--smooth) !important;
    }
    input:focus, textarea:focus {
        border-color: rgba(108,92,231,0.5) !important;
        box-shadow: 0 0 0 3px rgba(108,92,231,0.08), 0 0 24px rgba(108,92,231,0.1) !important;
        outline: none !important;
        background: rgba(255,255,255,0.05) !important;
    }
    section[data-testid="stExpander"] {
        background: rgba(255,255,255,0.02) !important;
        border: 1px solid rgba(255,255,255,0.05) !important;
        border-radius: var(--radius) !important;
        transition: all 0.3s var(--smooth) !important;
    }
    section[data-testid="stExpander"]:hover {
        border-color: rgba(108,92,231,0.2) !important;
    }
    .stTabs [data-baseweb="tab"] {
        color: var(--text-dim) !important; font-weight: 500 !important;
        transition: all 0.25s var(--smooth) !important;
    }
    .stTabs [data-baseweb="tab"]:hover { color: rgba(255,255,255,0.7) !important; }
    .stTabs [data-baseweb="tab"][aria-selected="true"] { color: var(--violet-ghost) !important; }
    .stTabs [data-baseweb="tab-list"] { border-bottom-color: rgba(255,255,255,0.04) !important; }
    .stTabs [data-baseweb="tab-highlight"] { background: var(--violet) !important; transition: all 0.3s var(--smooth) !important; }
    .stAlert { background: rgba(255,255,255,0.03) !important; border: 1px solid rgba(255,255,255,0.06) !important; border-radius: var(--radius) !important; }
    .stSpinner > div { border-top-color: var(--violet) !important; animation-duration: 0.6s !important; }
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-thumb { background: rgba(255,255,255,0.08); border-radius: 3px; }
    ::selection { background: rgba(108,92,231,0.4); color: #fff; }

    @keyframes fadeInUp {
        from { opacity: 0; transform: translateY(24px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @keyframes fadeInDown {
        from { opacity: 0; transform: translateY(-16px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @keyframes scaleIn {
        from { opacity: 0; transform: scale(0.94); }
        to { opacity: 1; transform: scale(1); }
    }

    @keyframes titleGlow {
        0%, 100% { filter: drop-shadow(0 0 8px rgba(162,155,254,0.3)); }
        50%      { filter: drop-shadow(0 0 24px rgba(108,92,231,0.5)); }
    }
    @keyframes gradientFlow {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    @keyframes cardReveal {
        from { opacity: 0; transform: translateY(20px) scale(0.97); }
        to   { opacity: 1; transform: translateY(0) scale(1); }
    }
    @keyframes subtleGlow {
        0%, 100% { box-shadow: 0 0 20px rgba(108,92,231,0.04); }
        50%      { box-shadow: 0 0 40px rgba(108,92,231,0.12); }
    }

    .main-title {
        font-size: 2.8em; font-weight: 700; text-align: center;
        background: linear-gradient(135deg, #C4B5FD 0%, #A29BFE 40%, #6C5CE7 100%);
        background-size: 200% 200%;
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        background-clip: text; margin-bottom: 0.1em; letter-spacing: 0.03em;
        animation: fadeInDown 0.7s var(--smooth) both,
                   gradientFlow 4s ease-in-out 0.7s infinite,
                   titleGlow 3s ease-in-out 1s infinite;
    }
    .sub-title {
        text-align: center; color: #8A8A9A; font-size: 1em; margin-bottom: 2.5em;
        letter-spacing: 0.05em; font-weight: 350;
        animation: fadeInUp 0.7s var(--smooth) 0.1s both;
    }
    .section-header {
        font-size: 1.1em; font-weight: 600; color: #C4B5FD;
        border-left: 3px solid rgba(108,92,231,0.5); padding-left: 14px;
        margin: 1.5em 0 0.6em 0;
        animation: fadeInUp 0.5s var(--smooth) both;
        transition: border-color 0.3s ease, color 0.3s ease;
    }
    .section-header:hover {
        border-left-color: rgba(162,155,254,0.8);
        color: #A29BFE;
    }
    .plan-box {
        background: rgba(255,255,255,0.025); backdrop-filter: blur(20px);
        -webkit-backdrop-filter: blur(20px);
        border: 1px solid rgba(255,255,255,0.06); border-radius: 14px;
        padding: 28px; color: #E8E6F0; line-height: 1.7;
        animation: cardReveal 0.55s var(--smooth) 0.15s both;
        transition: border-color 0.3s ease, box-shadow 0.3s ease, transform 0.3s var(--spring);
    }
    .plan-box:hover {
        border-color: rgba(108,92,231,0.15); box-shadow: 0 4px 28px rgba(108,92,231,0.06);
        transform: translateY(-1px);
    }
    .chat-bubble-user {
        background: rgba(108,92,231,0.08); backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(108,92,231,0.15); border-radius: 12px;
        padding: 12px 16px; margin: 8px 0; color: #E8E6F0;
        animation: cardReveal 0.4s var(--smooth) 0.08s both;
    }
    .chat-bubble-ai {
        background: rgba(162,155,254,0.06); backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(162,155,254,0.12); border-radius: 12px;
        padding: 12px 16px; margin: 8px 0; color: #E8E6F0;
        animation: cardReveal 0.4s var(--smooth) 0.12s both;
    }
    .pos-card {
        background: rgba(255,255,255,0.025); backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(255,255,255,0.05); border-radius: 14px;
        padding: 20px; text-align: center; color: #E8E6F0;
        transition: all 0.4s var(--spring);
        animation: cardReveal 0.5s var(--smooth) 0.2s both;
    }
    .pos-card:nth-child(1) { animation-delay: 0.20s; }
    .pos-card:nth-child(2) { animation-delay: 0.28s; }
    .pos-card:nth-child(3) { animation-delay: 0.36s; }
    .pos-card:nth-child(4) { animation-delay: 0.44s; }
    .pos-card:hover {
        background: rgba(255,255,255,0.05); border-color: rgba(108,92,231,0.3);
        box-shadow: 0 12px 40px rgba(108,92,231,0.12);
        transform: translateY(-4px) scale(1.02);
    }
    .pos-card:active { transform: scale(0.97); transition: all 0.1s ease; }
    .pos-card.selected {
        background: rgba(108,92,231,0.08); border-color: rgba(108,92,231,0.4);
        animation: subtleGlow 3s ease-in-out infinite;
    }

    /* ===== RESPONSIVE: MOBILE ===== */
    @media (max-width: 768px) {
        .main-title { font-size: 1.8em !important; }
        .sub-title { font-size: 0.85em !important; margin-bottom: 1.5em !important; }
        .section-header { font-size: 0.95em !important; }
        .plan-box { padding: 16px !important; }
        .stButton button { font-size: 0.82rem !important; }
        button[kind="primary"] { font-size: 0.85rem !important; }
        .pos-card { padding: 14px !important; }
        .chat-bubble-user, .chat-bubble-ai { font-size: 0.85em !important; }
        .stExpander { font-size: 0.85em !important; }
        section[data-testid="stSidebar"] { min-width: 260px !important; }
        footer { display: none !important; }
    }
    @media (max-width: 480px) {
        .main-title { font-size: 1.4em !important; }
        section[data-testid="stSidebar"] { width: 100vw !important; }
        .stTabs [data-baseweb="tab"] { font-size: 0.78em !important; padding: 6px 10px !important; }
        .stRadio label { font-size: 0.82em !important; }
        .stTextInput input, .stTextArea textarea { font-size: 16px !important; }
        button, .stButton button { min-height: 44px !important; }
    }
</style>
''', unsafe_allow_html=True)



def render_sidebar():
    with st.sidebar:
        st.markdown("### 🎤 辩论备赛助手")
        st.markdown("---")

        # 层级导航
        if st.session_state.page == "train":
            if st.button("← 返回备赛方案", use_container_width=True):
                st.session_state.page = "plan"
                st.rerun()
            if st.button("← 返回辩题输入", use_container_width=True):
                st.session_state.page = "input"
                st.session_state._splash_skipped = False
                st.rerun()
        elif st.session_state.page == "plan":
            if st.button("← 返回辩题输入", use_container_width=True):
                st.session_state.page = "input"
                st.session_state._splash_skipped = False
                st.rerun()

        if st.session_state.topic:
            st.markdown(f"**辩题：** {st.session_state.topic}")
            st.markdown(f"**持方：** {st.session_state.side}")
            if st.session_state.current_position:
                st.markdown(f"**辩位：** {st.session_state.current_position}")
            # 保存按钮
            if st.button("💾 保存当前进度", use_container_width=True, help="保存后下次可继续备赛"):
                save_session()
                st.success("进度已保存 ✓")
        st.markdown("---")
        _render_model_config()
        st.markdown("---")
        st.markdown("*由 Claude 提供 AI 能力*")

        # 论据库通知
        if st.session_state.evidence_notification:
            st.markdown("---")
            st.info(st.session_state.evidence_notification)
            if st.button("已读", key="sidebar_clear_notification"):
                st.session_state.evidence_notification = ""
                st.rerun()


def _render_model_config():
    llm_key = read_env("LLM_API_KEY") or read_env("ANTHROPIC_API_KEY")
    exa_key = read_env("EXA_API_KEY")
    has_valid_key = llm_key and len(llm_key) > 20 and "your_api_key" not in llm_key.lower()

    # 检测是否运行在 Streamlit Cloud（secrets 中读取到值 = 云端配置）
    try:
        import streamlit as st
        is_cloud = hasattr(st, "secrets") and st.secrets and ("LLM_API_KEY" in st.secrets or "ANTHROPIC_API_KEY" in st.secrets)
    except Exception:
        is_cloud = False

    with st.expander("⚙️ 模型配置" if has_valid_key else "⚙️ 模型配置（首次使用请点这里）", expanded=not has_valid_key):
        if is_cloud:
            st.success(f"✅ AI 模型已由管理员配置（{read_env('LLM_MODEL', 'deepseek-chat')}），无需额外设置，可直接使用。")
            st.caption("如需修改，请自行本地部署并配置自己的 API Key。")
            return

        provider = st.selectbox(
            "接口类型",
            options=["anthropic", "openai"],
            index=0 if read_env("LLM_PROVIDER", "anthropic") == "anthropic" else 1,
            format_func=lambda x: "Anthropic 原生" if x == "anthropic" else "OpenAI 兼容",
            key="cfg_provider",
        )

        api_key = st.text_input("API Key", value=llm_key, type="password",
                                placeholder="sk-...", key="cfg_api_key")

        if provider == "openai":
            base_url = st.text_input("接口地址", value=read_env("LLM_BASE_URL", "https://api.anthropic.com/v1"),
                                     placeholder="https://api.openai.com/v1", key="cfg_base_url")
        else:
            base_url = "https://api.anthropic.com/v1"

        model_opts = {
            "anthropic": ["claude-sonnet-4-6", "claude-opus-4-7", "claude-haiku-4-5", "claude-sonnet-4-5"],
            "openai": ["deepseek-chat", "deepseek-reasoner", "gpt-4o", "gpt-4-turbo", "qwen-plus", "qwen-max"],
        }
        opts = model_opts.get(provider, ["claude-sonnet-4-6"])
        cur_model = read_env("LLM_MODEL", "claude-sonnet-4-6")
        if cur_model not in opts:
            opts.insert(0, cur_model)
        try:
            idx = opts.index(cur_model)
        except ValueError:
            idx = 0
        model = st.selectbox("模型", options=opts, index=idx, key="cfg_model")

        use_custom = st.checkbox("自定义模型名", key="cfg_custom_model")
        if use_custom:
            model = st.text_input("模型名称", value=cur_model, key="cfg_custom_model_input")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 保存配置", use_container_width=True):
                try:
                    write_env("LLM_PROVIDER", provider)
                    write_env("LLM_API_KEY", api_key)
                    write_env("LLM_BASE_URL", base_url)
                    write_env("LLM_MODEL", model)
                    reset_llm()
                    st.success("配置已保存 ✓ 现在可以输入辩题开始备赛了！")
                except Exception as e:
                    st.error(f"保存失败：{e}")
        with col2:
            if st.button("🔌 测试连接", use_container_width=True, disabled=not api_key):
                with st.spinner("测试中..."):
                    try:
                        t = get_llm({"provider": provider, "api_key": api_key,
                                     "base_url": base_url, "model": model})
                        r = t.ask("你是测试助手", "请回复'连接成功'")
                        if len(r) > 0:
                            st.success("连接成功 ✓")
                        else:
                            st.error(f"响应异常：{r[:100]}")
                    except Exception as e:
                        st.error(f"连接失败：{str(e)[:200]}")

        s_llm = "✅ 已配置" if has_valid_key else "⚠️ 未配置"
        s_exa = "✅ 已配置" if exa_key else "⚠️ 未配置"
        st.caption(f"LLM：{s_llm} | Exa：{s_exa}")

        # 启动画面设置
        show_splash = read_env("SHOW_SPLASH", "true").lower() == "true"
        new_splash = st.checkbox("启动时显示开场动画", value=show_splash, key="cfg_splash",
                                 help="每次打开程序时是否先展示「辩」字开场画面")
        if new_splash != show_splash:
            write_env("SHOW_SPLASH", "true" if new_splash else "false")
            st.rerun()


# ---- 页面1：辩题输入 ----
def render_input_page():
    # ===== 启动画面（原生Streamlit方案，100%可靠） =====
    show_splash = read_env("SHOW_SPLASH", "true").lower() == "true"
    if show_splash and not st.session_state.get("_splash_skipped"):
        from splash_b64 import SPLASH_BIAN_B64
        st.markdown(f"""
        <style>
            @keyframes fadeUpSplash {{
                from {{ opacity: 0; transform: translateY(40px); }}
                to   {{ opacity: 1; transform: translateY(0); }}
            }}
            @keyframes glowPulseSplash {{
                0%, 100% {{ filter: drop-shadow(0 0 24px rgba(108,92,231,0.35)) drop-shadow(0 0 8px rgba(201,169,110,0.2)); }}
                50%      {{ filter: drop-shadow(0 0 56px rgba(162,155,254,0.6)) drop-shadow(0 0 16px rgba(201,169,110,0.35)); }}
            }}
            @keyframes breatheSplash {{
                0%, 100% {{ opacity: 0.3; }}
                50%      {{ opacity: 1; }}
            }}
            @keyframes inkFadeIn {{
                0%   {{ opacity: 0; filter: blur(8px); transform: scale(1.06); }}
                60%  {{ opacity: 0.9; filter: blur(1px); transform: scale(0.98); }}
                100% {{ opacity: 1; filter: blur(0); transform: scale(1); }}
            }}
            .splash-hero {{
                text-align: center; padding-top: 8vh;
                animation: fadeUpSplash 0.8s cubic-bezier(0.25, 0.1, 0.25, 1);
            }}
            .splash-img {{
                width: 320px; max-width: 50vw;
                animation: inkFadeIn 1.2s cubic-bezier(0.25, 0.1, 0.25, 1) both,
                           glowPulseSplash 3.5s ease-in-out 1.5s infinite;
                margin-bottom: 0.3em;
            }}
            .splash-subtitle {{
                color: rgba(255,255,255,0.32); font-size: 1.05em;
                letter-spacing: 0.4em; margin-bottom: 3em;
                animation: fadeUpSplash 0.9s cubic-bezier(0.25, 0.1, 0.25, 1) 0.2s both;
            }}
            .splash-hint-text {{
                color: rgba(255,255,255,0.2); font-size: 0.85em;
                animation: breatheSplash 2.5s ease-in-out infinite;
            }}
        </style>
        <div class="splash-hero">
            <img class="splash-img" src="data:image/png;base64,{SPLASH_BIAN_B64}" alt="辩">
            <div class="splash-subtitle">DEBATE  ASSISTANT</div>
        </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            if st.button("开始备赛", type="primary", use_container_width=True, key="splash_enter"):
                st.session_state._splash_skipped = True
                st.rerun()
        # 提示文字放在按钮下方
        _, c2, _ = st.columns([1, 1, 1])
        with c2:
            st.markdown('<p class="splash-hint-text" style="text-align:center;">轻触上方按钮进入</p>', unsafe_allow_html=True)
        st.stop()

    st.markdown('<h1 class="main-title">辩论备赛助手</h1>', unsafe_allow_html=True)
    st.markdown('<p class="sub-title">一站式辩论备赛工具 —— 资料搜索 · 方案生成 · 分辩位训练</p>', unsafe_allow_html=True)

    llm_key = read_env("LLM_API_KEY") or read_env("ANTHROPIC_API_KEY")
    has_valid_key = llm_key and len(llm_key) > 20 and "your_api_key" not in llm_key.lower()
    if not has_valid_key:
        st.warning("🔑 请先在侧边栏「模型配置」中输入你的 API Key 并保存。支持 Anthropic / DeepSeek / OpenAI 等接口。", icon="🔑")

    # ===== 继续备赛：已保存的会话 =====
    saved = list_saved_sessions()
    if saved:
        with st.expander(f"📂 继续之前的备赛（{len(saved)} 个已保存的会话）", expanded=True):
            for i, s in enumerate(saved):
                cols = st.columns([4, 2, 1.5, 1])
                with cols[0]:
                    st.markdown(f"**{s['topic']}**")
                    status_parts = [f"{s['side']}方"]
                    if s["has_plan"]:
                        status_parts.append("有备赛方案")
                    if s["has_standpoint"]:
                        status_parts.append("立论已确认")
                    st.caption(" · ".join(status_parts))
                with cols[1]:
                    st.caption(f"保存于 {s['saved_at']}")
                with cols[2]:
                    if st.button(f"▶️ 继续", key=f"resume_{i}", use_container_width=True):
                        load_session(s["topic"])
                        st.success(f"已恢复「{s['topic']}」的备赛进度！")
                        st.rerun()
                with cols[3]:
                    if st.button(f"🗑️", key=f"del_save_{i}", help="删除此存档"):
                        delete_session(s["topic"])
                        st.rerun()

    # ===== 新建备赛 =====
    st.markdown("---")
    st.markdown("### 🆕 新建备赛")
    col1, col2 = st.columns([3, 1])
    with col1:
        topic = st.text_input("📋 请输入辩题", value=st.session_state.topic,
                              placeholder="例如：人工智能应该取代人类决策吗？")
    with col2:
        side = st.selectbox("🏷️ 你的持方", ["正方", "反方"],
                            index=0 if st.session_state.side == "正方" else 1)

    st.session_state.topic = topic
    st.session_state.side = side

    with st.expander("🎬 视频链接（可选）— 提取过往比赛字幕"):
        v = st.text_area("输入 YouTube 或 B站 视频链接，每行一个",
                         value=st.session_state.video_url_str,
                         placeholder="https://www.youtube.com/watch?v=xxx\nhttps://www.bilibili.com/video/BVxxx")
        st.session_state.video_url_str = v

    st.markdown("---")
    if st.button("🔍 开始搜索资料 & 生成备赛方案", type="primary",
                 use_container_width=True, disabled=not topic):
        if not topic:
            st.warning("请先输入辩题")
            return

        v_urls = [u.strip() for u in st.session_state.video_url_str.split("\n") if u.strip()]

        with st.spinner("⏳ 正在搜索过往比赛资料、哲学讨论、热点新闻..."):
            st.session_state.research_result = do_research(topic, side, v_urls or None)

        with st.spinner("📝 正在生成备赛方案..."):
            st.session_state.plan = generate_plan(topic, side, st.session_state.research_result)

        with st.spinner("💡 正在生成辅助论点..."):
            st.session_state.arguments = generate_arguments(topic, side, st.session_state.research_result)

        with st.spinner("🧱 正在生成立论积木块..."):
            generate_blocks(topic, side, st.session_state.research_result)

        st.session_state.standpoint_confirmed = False
        st.session_state.user_arguments = []
        st.session_state.page = "plan"
        save_session()  # 自动保存
        st.rerun()


# ---- 页面2：方案查看 & 立论构建 & 辩位选择 ----
def render_plan_page():
    st.markdown('<h1 class="main-title">备赛方案</h1>', unsafe_allow_html=True)
    st.markdown(f'<p class="sub-title">{st.session_state.topic} · {st.session_state.side}方</p>', unsafe_allow_html=True)

    t1, t2, t3 = st.tabs(["📋 备赛方案", "💡 辅助论点", "📚 原始资料"])
    with t1:
        st.markdown(st.session_state.plan)
        st.download_button("📥 导出备赛方案为 Word",
                           data=generate_docx("备赛方案", st.session_state.plan),
                           file_name=f"备赛方案_{st.session_state.topic}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with t2:
        st.markdown(st.session_state.arguments)
    with t3:
        with st.expander("查看搜索到的原始资料"):
            st.text(st.session_state.research_result[:30000])

    # ===== 积木式立论构建区（核心） =====
    st.markdown("---")
    st.markdown('<h2 class="section-header">🧱 积木式立论构建（核心任务）</h2>', unsafe_allow_html=True)
    st.caption("像拼积木一样搭建你的立论框架——勾选你想要的判准、定义、论点和逻辑链，组合成完整的立论体系。")

    if st.session_state.standpoint_confirmed:
        st.success("✅ 立论框架已确认 —— 可进入分辩位训练")
        with st.expander("📋 查看已确认的立论框架", expanded=False):
            _render_assembled_framework()
        if st.button("🔄 重新搭建立论", use_container_width=True):
            _reset_blocks()
            st.rerun()
    else:
        st.info("🧱 在下方四类积木中勾选你想要的模块，也可以添加自定义积木。满意后点击「确认立论框架」。")

        if not st.session_state.blocks_generated:
            st.warning("积木块尚未生成，请返回辩题输入页重新搜索。")
        else:
            _render_block_picker()
            _render_custom_block_input()
            _render_framework_preview()

            st.markdown("---")
            if st.button("🔒 确认立论框架", type="primary", use_container_width=True):
                _assemble_framework()
                review_framework()
                save_session()
                st.rerun()

    # ===== 框架审查结果（确认后显示） =====
    if st.session_state.standpoint_confirmed and st.session_state.framework_reviewed:
        st.markdown("---")
        st.markdown('<h2 class="section-header">🔍 框架审查与辩位任务</h2>', unsafe_allow_html=True)

        with st.expander("📋 框架审查意见", expanded=True):
            st.markdown(st.session_state.framework_review)

        with st.expander("🎯 各辩位任务分配（新国辩赛制）", expanded=True):
            st.markdown(st.session_state.position_tasks.get("full", ""))
        st.markdown("---")

    # ===== 论据库 =====
    _render_evidence_panel()

    # ===== 辩位选择（立论确认后显示） =====
    if st.session_state.standpoint_confirmed:
        st.markdown('<h2 class="section-header">选择辩位开始针对性训练</h2>', unsafe_allow_html=True)
        positions = [
            ("一辩", "立论陈词 · 前端接质", "一辩陈词搭建框架；前端接质询守护定义与判准"),
            ("二辩", "申论驳论 · 中端接质", "回应遗留问题+申论深化+驳论攻击；中端战场核心接质"),
            ("三辩", "质询盘问 · 中场收束", "质询对方中端核心战场；中场总结与论点概括收束"),
            ("四辩", "全局质询 · 价值收束", "质询对方前端问题；全局战场总结与价值升华"),
        ]
        cols = st.columns(4)
        for i, (pos, role, duty) in enumerate(positions):
            with cols[i]:
                selected = st.session_state.current_position == pos
                cls = "pos-card selected" if selected else "pos-card"
                st.markdown(f"""
                <div class="{cls}"><h3>{pos}</h3>
                <p style="color:#667eea;font-weight:600;">{role}</p>
                <p style="color:#888;font-size:0.85em;">{duty}</p></div>
                """, unsafe_allow_html=True)
                if st.button(f"🎯 {pos}备赛", key=f"sel_{pos}", use_container_width=True):
                    st.session_state.current_position = pos
                    st.session_state.chat_history = []
                    st.session_state.cross_exam_rounds = []
                    st.session_state.cross_exam_current_q = ""
                    st.session_state.cross_exam_evaluation = ""
                    st.session_state.page = "train"
                    st.rerun()
    else:
        st.markdown("---")
        st.warning("👆 请先完成上方「积木式立论构建」并点击「确认立论框架」，再进入分辩位训练。")


# ---- 积木式立论辅助函数 ----
def _render_block_picker():
    """渲染四类积木选择器"""
    categories = [
        ("🚪 开题思路", "approach", st.session_state.approach_blocks, "选择辩论的破题角度和论证切入方式"),
        ("⚖️ 判准", "criteria", st.session_state.criteria_blocks, "选择用来评判辩论胜负的根本标准"),
        ("📖 定义", "definitions", st.session_state.definition_blocks, "选择辩题关键概念的解读方式"),
        ("💡 论点", "arguments", st.session_state.argument_blocks, "选择要使用的核心论点"),
        ("🔗 逻辑链", "logic_chains", st.session_state.logic_blocks, "选择从前提推导到结论的推理链条"),
    ]

    for cat_name, cat_key, blocks, help_text in categories:
        if not blocks:
            continue
        selected_count = sum(1 for b in blocks if b["selected"])
        with st.expander(f"{cat_name} （{selected_count}/{len(blocks)} 已选）", expanded=True):
            st.caption(help_text)
            for i, block in enumerate(blocks):
                cols = st.columns([0.05, 0.85, 0.1])
                with cols[0]:
                    checked = st.checkbox("", value=block["selected"], key=f"chk_{cat_key}_{i}",
                                          label_visibility="collapsed")
                    if checked != block["selected"]:
                        block["selected"] = checked
                        st.rerun()
                with cols[1]:
                    bg = "#f0f8ff" if block["selected"] else "#fafafa"
                    border = "2px solid #667eea" if block["selected"] else "1px solid #ddd"
                    st.markdown(f"""
                    <div style="background:{bg}; color:#1a1a1a; border:{border}; border-radius:8px; padding:12px; margin:2px 0;">
                        {block['text']}
                    </div>
                    """, unsafe_allow_html=True)
                with cols[2]:
                    if st.button("🔄", key=f"regen_{cat_key}_{i}", help="重新生成此积木"):
                        _regenerate_single_block(cat_key, i)
                        st.rerun()


def _render_custom_block_input():
    """自定义积木输入"""
    st.markdown("---")
    st.markdown("### ✨ 添加自定义积木")
    col_type, col_input = st.columns([1, 3])
    with col_type:
        cat = st.selectbox("积木类型", ["论点", "判准", "定义", "逻辑链"], key="custom_cat")
    with col_input:
        idea = st.text_input("输入你的想法或论点",
                             placeholder="例如：从隐私权的角度，AI决策对个人数据的收集本身就构成侵权...",
                             key="custom_idea")
    col_add, col_expand = st.columns(2)
    with col_add:
        if st.button("➕ 添加自定义积木", use_container_width=True, disabled=not idea):
            st.session_state.custom_blocks.append({"text": idea, "category": cat})
            st.success(f"已添加自定义{cat}！")
            st.rerun()
    with col_expand:
        if st.button("🤖 AI 拓展并添加", use_container_width=True, disabled=not idea):
            with st.spinner("AI 正在拓展..."):
                llm = get_llm()
                expanded = llm.ask(
                    f"你是辩论教练。请将用户的{cat}草稿拓展为精炼的辩论积木块（60-200字），含主张+依据+逻辑。只输出拓展后的文本，不要其他。",
                    f"辩题：{st.session_state.topic}\n持方：{st.session_state.side}\n用户草稿：{idea}",
                    temperature=0.8)
                st.session_state.custom_blocks.append({"text": f"【AI拓展】{expanded.strip()}", "category": cat})
                st.success(f"AI 已拓展并添加为{cat}！")
                st.rerun()

    # 显示已添加的自定义积木
    if st.session_state.custom_blocks:
        st.markdown("**已添加的自定义积木：**")
        for i, cb in enumerate(st.session_state.custom_blocks):
            cols = st.columns([0.1, 0.8, 0.1])
            with cols[1]:
                st.markdown(f"**[{cb['category']}]** {cb['text'][:200]}")
            with cols[2]:
                if st.button("🗑️", key=f"del_custom_{i}"):
                    st.session_state.custom_blocks.pop(i)
                    st.rerun()


def _render_framework_preview():
    """实时预览选中的立论框架"""
    selected_approaches = [b for b in st.session_state.approach_blocks if b["selected"]]
    selected_criteria = [b for b in st.session_state.criteria_blocks if b["selected"]]
    selected_defs = [b for b in st.session_state.definition_blocks if b["selected"]]
    selected_args = [b for b in st.session_state.argument_blocks if b["selected"]]
    selected_logics = [b for b in st.session_state.logic_blocks if b["selected"]]
    custom = st.session_state.custom_blocks

    total = len(selected_approaches) + len(selected_criteria) + len(selected_defs) + len(selected_args) + len(selected_logics) + len(custom)
    if total == 0:
        return

    st.markdown("---")
    st.markdown(f"### 🔍 立论预览（已选 {total} 个积木块）")

    if selected_approaches:
        st.markdown("**🚪 开题思路：**")
        for b in selected_approaches:
            st.markdown(f"- {b['text'][:150]}")

    if selected_criteria:
        st.markdown("**⚖️ 判准：**")
        for b in selected_criteria:
            st.markdown(f"- {b['text'][:120]}")

    if selected_defs:
        st.markdown("**📖 定义：**")
        for b in selected_defs:
            st.markdown(f"- {b['text'][:120]}")

    if selected_args:
        st.markdown("**💡 论点：**")
        for b in selected_args:
            st.markdown(f"- {b['text'][:120]}")

    if selected_logics:
        st.markdown("**🔗 逻辑链：**")
        for b in selected_logics:
            st.markdown(f"- {b['text'][:120]}")

    if custom:
        st.markdown("**✨ 自定义：**")
        for b in custom:
            st.markdown(f"- [{b['category']}] {b['text'][:120]}")


def _assemble_framework():
    """组装立论框架，写入 st.session_state.arguments"""
    parts = []

    parts.append("# 立论框架\n")

    selected_approaches = [b for b in st.session_state.approach_blocks if b["selected"]]
    if selected_approaches:
        parts.append("## 🚪 开题思路")
        parts.extend(f"- {b['text']}" for b in selected_approaches)

    selected_criteria = [b for b in st.session_state.criteria_blocks if b["selected"]]
    if selected_criteria:
        parts.append("## ⚖️ 判准")
        parts.extend(f"- {b['text']}" for b in selected_criteria)

    selected_defs = [b for b in st.session_state.definition_blocks if b["selected"]]
    if selected_defs:
        parts.append("## 📖 核心定义")
        parts.extend(f"- {b['text']}" for b in selected_defs)

    selected_args = [b for b in st.session_state.argument_blocks if b["selected"]]
    if selected_args:
        parts.append("## 💡 核心论点")
        parts.extend(f"### 论点 {i+1}\n{b['text']}" for i, b in enumerate(selected_args))

    selected_logics = [b for b in st.session_state.logic_blocks if b["selected"]]
    if selected_logics:
        parts.append("## 🔗 逻辑链条")
        parts.extend(f"- {b['text']}" for b in selected_logics)

    custom = st.session_state.custom_blocks
    if custom:
        parts.append("## ✨ 自定义补充")
        parts.extend(f"- [{b['category']}] {b['text']}" for b in custom)

    st.session_state.arguments = "\n\n".join(parts)
    st.session_state.standpoint_confirmed = True


def _render_assembled_framework():
    """查看已确认的立论框架"""
    st.markdown(st.session_state.arguments)


def _reset_blocks():
    """重置所有积木和立论状态"""
    for blocks in [st.session_state.approach_blocks, st.session_state.criteria_blocks,
                    st.session_state.definition_blocks, st.session_state.argument_blocks,
                    st.session_state.logic_blocks]:
        for b in blocks:
            b["selected"] = False
    st.session_state.custom_blocks = []
    st.session_state.standpoint_confirmed = False
    st.session_state.framework_reviewed = False
    st.session_state.framework_review = ""
    st.session_state.position_tasks = {}
    st.session_state.position_speeches = {}
    st.session_state.user_arguments = []


def _regenerate_single_block(cat_key: str, index: int):
    """重新生成单个积木块"""
    llm = get_llm()
    cat_names = {"approach": "开题思路", "criteria": "判准", "definitions": "定义", "arguments": "论点", "logic_chains": "逻辑链"}
    cat_name = cat_names.get(cat_key, cat_key)
    # cat_key -> session state attribute mapping
    attr_map = {"approach": "approach_blocks", "criteria": "criteria_blocks",
                "definitions": "definition_blocks", "arguments": "argument_blocks",
                "logic_chains": "logic_blocks"}
    attr = attr_map.get(cat_key, f"{cat_key}_blocks")
    block_list = getattr(st.session_state, attr)
    old_text = block_list[index]["text"] if index < len(block_list) else ""

    prompt = f"""请为辩题「{st.session_state.topic}」（持方：{st.session_state.side}）生成一个新的{cat_name}积木块。
要求：60-200字，精炼独立可用，与以下已有积木不重复。
已有积木：{old_text[:100]}...

请直接输出新的{cat_name}文本，不要其他文字。"""
    new_text = llm.ask("你是辩论立论架构师。生成精炼的辩论积木块。", prompt, temperature=0.9)

    if new_text and len(new_text) > 10:
        block_list[index] = {"text": new_text.strip(), "selected": block_list[index]["selected"]}


def review_framework():
    """审查立论框架的逻辑一致性，生成优化建议和辩位任务"""
    llm = get_llm()
    framework = st.session_state.arguments

    # Step 1: 审查框架
    with st.spinner("🔍 正在审查框架逻辑一致性..."):
        review = llm.ask(
            """你是资深辩论裁判和立论教练。请审查以下立论框架的逻辑一致性。

审查要点：
1. 各积木块之间是否存在逻辑矛盾？（如判准与定义冲突、论点与逻辑链脱节）
2. 开题思路是否与后续论证方向一致？
3. 判准和定义是否支持所有选定的论点？
4. 逻辑链是否能有效串联所有论点？

如果有矛盾，请明确指出冲突点，给出修改建议，并解释为什么修改后更好。
如果逻辑一致，请进行完善优化，让表述更精准、论证更紧密。
最后给出优化后的完整框架。""",
            f"辩题：{st.session_state.topic}\n持方：{st.session_state.side}\n\n当前立论框架：\n{framework}",
            temperature=0.7)
        st.session_state.framework_review = review

    # Step 2: 生成各辩位任务
    with st.spinner("📋 正在制定各辩位任务..."):
        tasks = llm.ask(
            """你是辩论教练。请基于以下立论框架，为新国辩赛制制定每个辩位的具体任务。

新国辩赛制说明（按比赛时序）：
- 一辩：立论陈词（3分钟）+ 前端接质询（被对方四辩质询）→ 搭建框架，定义概念，守护定义与判准
- 二辩：申论驳论陈词（2分钟）+ 对辩（1.5分钟）+ 中端接质询（被对方三辩质询）→ 回应遗留+申论深化+驳论攻击，守护中端战场
- 三辩：盘问质询对方中端战场（2.5分钟）+ 中场小结收束（1.5分钟）→ 主动出击，设计问题链攻击对方中端论证，最后收束战场
- 四辩：质询对方前端问题（先质询）+ 全局总结陈词（3.5分钟）→ 攻击对方定义/判准/前提，全局战场总结+价值升华

输出格式：
【一辩任务】
- 立论陈词核心任务：...
- 前端接质防守要点（对方四辩会攻击的）：...
- 需要预设的定义防守：...

【二辩任务】
- 申论驳论核心任务：...
- 需要回应的对方一辩论点：...
- 中端接质防守要点（对方三辩会攻击的）：...
- 对辩策略：...

【三辩任务】
- 质询对方中端战场的核心方向（3-5组问题链）：...
- 对方中端论证的薄弱点：...
- 中场收束小结的概括策略：...

【四辩任务】
- 质询对方前端问题的方向（定义/判准/前提）：...
- 全局战场总结要点：...
- 价值升华方向与收尾金句：...""",
            f"辩题：{st.session_state.topic}\n持方：{st.session_state.side}\n\n立论框架：\n{framework}\n\n审查意见：\n{st.session_state.framework_review[:3000]}",
            temperature=0.8)
        st.session_state.position_tasks = {"full": tasks}

    st.session_state.framework_reviewed = True


def generate_position_speech(position: str):
    """基于框架和辩位任务，生成该辩位的初稿辩词（新国辩赛制）"""
    llm = get_llm()
    profile = POSITION_PROFILES.get(position, POSITION_PROFILES["一辩"])
    framework = st.session_state.arguments
    review = st.session_state.framework_review
    tasks = st.session_state.position_tasks.get("full", "")

    prompt = f"""请基于以下立论框架，为{position}（{profile['role']}）撰写新国辩赛制的初稿辩词。

辩题：{st.session_state.topic}
持方：{st.session_state.side}

立论框架：
{framework}

框架审查意见：
{review[:2000]}

辩位任务：
{tasks[:3000]}

请严格按照{position}的辩词结构和风格要求（参考教练提示），生成完整的初稿辩词。
直接输出可使用的辩词文本，包含具体的论证内容、数据和例子。"""

    return llm.ask(profile["system_prompt"], prompt, temperature=0.8)


# ============================================================
# 8.5 接质训练模块
# ============================================================
def generate_cross_question(position: str) -> str:
    """基于立论框架和辩位特点生成针对性质询问题"""
    llm = get_llm()
    framework = st.session_state.arguments
    history_text = ""
    if st.session_state.cross_exam_rounds:
        history_text = "此前质询记录：\n" + "\n".join(
            f"Q: {r['question']}\nA: {r.get('user_answer') or r.get('ai_answer','')}"
            for r in st.session_state.cross_exam_rounds[-3:]
        )

    # 各辩位质询策略（AI扮演对方对应辩位，模拟比赛中的实际质询方）
    strategies = {
        "一辩": """你扮演对方四辩，正在质询我方一辩（前端接质）。
我方一辩刚完成立论陈词，你作为对方四辩的攻击重点是前端问题——定义、判准、前提。

质询重点（按优先级）：
1. 攻定义自洽性：我方的核心定义是否内部一致？是否有模糊可争夺的空间？
2. 挑判准偏差：我方的判准是否偏向我方立场？能否覆盖辩题的全部情况？
3. 挖前提反例：我方立论中最薄弱的隐含前提是什么？能否找到反例？
4. 逼定义让步：通过封闭式问题逼我方一辩在定义上做出不利收缩
问题风格：简短封闭式（是/否），20-30字，模拟四辩的前端攻击节奏。""",

        "二辩": """你扮演对方三辩，正在质询我方二辩（中端接质）。
此时对辩环节已展开，你作为对方三辩的攻击重点是我方二辩的申论深化和例证。

质询重点（按优先级）：
1. 设情景陷阱：构造具体案例/情景，逼我方二辩在该情景下难以自圆其说
2. 挖例证漏洞：我方二辩引用的例子是否具有普遍性？数据来源是否权威？
3. 推极端后果：如果我方二辩的逻辑成立，会导致什么荒谬推论？
4. 找逻辑断层：我方二辩在申论深化中哪个推理环节最脆弱？
问题风格：先设情景再追问（"假如...那你方是否认为..."），30-50字。""",

        "三辩": """我方三辩是主动质询方（质询对方中端核心战场）。此训练中你扮演对方二辩反制我方的质询。
你模拟对方二辩被我方三辩质询时的反制：

质询重点：
1. 反制归谬策略：预判我方三辩可能使用的归谬法/类比法，提前堵截
2. 预设拆解：指出我方三辩问题链背后的预设错误
3. 情景反例：给出一个反例挑战我方三辩的问题链前提
问题风格：短促有力，15-25字，打乱我方三辩的质询节奏。""",

        "四辩": """我方四辩是主动质询方（质询对方前端问题）。此训练中你扮演对方一辩反制我方的质询。
你模拟对方一辩被我方四辩攻击前端时的回应：

质询重点：
1. 坚守定义：对方四辩攻击我方框架根基，你模拟对方一辩如何防守
2. 价值对冲：从对方的价值预设中找到可以反制的点
3. 逼压边界：逼我方四辩在质询时间压力下暴露追问的预设问题
问题风格：防守反击型，20-30字。""",
    }

    strategy = strategies.get(position, strategies["一辩"])
    round_count = len(st.session_state.cross_exam_rounds)
    phase_hint = ""
    if round_count == 0:
        phase_hint = "这是第一轮质询，请从策略中的第1点入手。"
    elif round_count <= 2:
        phase_hint = f"这是第{round_count+1}轮质询，请根据对方上一轮的回应，递进追问或转向策略中的第{min(round_count+1,4)}点。"
    else:
        phase_hint = f"这是第{round_count+1}轮质询，请综合运用策略，寻找对方回应中的矛盾点进行追问。"

    prompt = f"""你正在模拟新国辩赛制中的质询环节。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
我方立论框架：
{framework[:3000]}

{history_text}

{strategy}

{phase_hint}

请提出一个质询问题。直接输出问题，不要其他文字。"""

    return llm.ask("你是对方辩手，正在质询我方。严格按辩位特点和策略提问。只输出问题。", prompt, temperature=0.9)


def generate_cross_response(position: str, question: str) -> str:
    """为接质方生成回应（用于学习参考）"""
    llm = get_llm()
    framework = st.session_state.arguments

    # 各辩位接质风格
    response_styles = {
        "一辩": """你现在是我方一辩（立论陈词·前端接质），正在接对方四辩的质询。接质要点：
- 你是框架的守护者：对方四辩专门攻击你的定义和判准，任何回答都不能松动核心概念
- 遇到定义类质疑：坚持我方定义，阐明其合理性和必要性，不给对方概念争夺的空间
- 遇到判准性质疑：指出我方判准的公平性和覆盖面
- 遇到前提反例：指明反例的特殊性，不构成对我方前提的否定
- 风格：温和而坚定，不失风度，前端接质是对方四辩的首要攻击目标""",
        "二辩": """你现在是我方二辩（申论驳论·中端接质），正在接对方三辩的质询。接质要点：
- 你是攻防核心：对方三辩专门攻击你中端的论证和例证，你需要灵活应对
- 面对情景陷阱：指出情景的特殊性，或接受情景但扭回有利于我方的解释
- 遇到例证质疑：承认单一例证的局限，但强调我方多条例证形成的证据链
- 遇到逻辑追问：理清逻辑链的每个环节，指出对方误解了我方的推理关系
- 风格：灵活机敏，善于化解，守护中端战场论证的可靠性""",
        "三辩": """你现在是我方三辩（质询盘问·中场收束），正在应对对方二辩的反制。接质要点：
- 你是我方主动质询方，面对反制时要冷静：不能被对方打乱质询节奏
- 遇到反制性预设拆解：立即指出对方反制本身的逻辑错误
- 遇到情景反例：指出反例不适用，或反例实质上支持我方立场
- 保持质询主动权：简短有力回应，迅速回到自己的问题链
- 风格：沉稳冷静，不被反制带偏，保持中场收束的主动权""",
        "四辩": """你现在是我方四辩（全局质询·价值收束），正在应对对方一辩的反制。接质要点：
- 你是全局的定音者：每个回应都是价值层面的小总结
- 遇到对方一辩的坚守：不纠缠具体定义细节，转而指出对方一辩设定中隐含的价值偏向
- 遇到价值对冲：用更包容的价值框架吸收对方的合理关切
- 将反制转化为升华机会：每次回应都为后续的总结陈词积累价值素材
- 风格：大局观强，不被拖入细节纠缠，始终站在全局高度""",
    }

    style = response_styles.get(position, response_styles["一辩"])

    prompt = f"""你是辩论赛中{position}（{POSITION_PROFILES.get(position, {}).get('role','')}），正在接对方质询。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
我方立论框架：
{framework[:3000]}

对方质询问题：{question}

{style}

⚠️ 铁律：接质询位绝对不能反问对方！不能用问题回答问题。必须正面回应，陈述观点。

请给出简洁有力的回应（40-100字）。直接输出回应文本。"""

    return llm.ask("你是新国辩赛制中的接质辩手。正面回应、不反问。", prompt, temperature=0.8)


def evaluate_cross_answer(position: str, question: str, answer: str) -> str:
    """评价用户的接质回应，给出优化建议，并建议AI对手的下一步"""
    llm = get_llm()
    framework = st.session_state.arguments

    # 一辩/二辩接质规则提醒
    no_askback = ""
    if position in ("一辩", "二辩"):
        no_askback = f"\n⚠️ 重要规则：{position}是被质询方，没有追问对方的权力，只能正面接质回应。评价时注意：如果回应中有反问或提问，这是违规的。"

    prompt = f"""你是资深辩论教练，正在评价我方{position}的接质回应。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
立论框架：
{framework[:3000]}

对方质询问题：{question}
我方{position}的回应：{answer}
{no_askback}

请从以下维度评价并给出建议：

**评分**（满分10分）：
- 防守力（是否守住我方立场）：X/4
- 合规性（是否正面回应、有无违规反问）：X/3
- 技巧性（用词精准度、有无化解陷阱转化优势）：X/3
- 总分：X/10

**对你的优化建议**（如何更好地接质回应）：
- 做得好：
- 可改进：
- 练习要点：

**对AI对手下一步的建议**（这不是给你的建议，是给AI对手的）：
- AI应继续追问 or AI应更换话题？
- 理由：

请简洁输出，控制在200字以内。直接给出评价。"""

    return llm.ask("你是资深辩论教练。客观评价接质回应。注意一辩二辩不能追问对方。", prompt, temperature=0.7)


def decide_next_question(position: str, evaluation: str) -> str:
    """根据评价决定下一问：追问 or 换题"""
    llm = get_llm()
    framework = st.session_state.arguments

    # 各辩位的追问方向指引
    followup_guides = {
        "一辩": "你扮演对方四辩，追问继续扣我方一辩前端的定义/判准/前提，递进一层施压。",
        "二辩": "你扮演对方三辩，追问沿着我方二辩中端回应的逻辑延伸，找新矛盾或延伸情景。",
        "三辩": "你扮演对方二辩反制，追问针对我方三辩问题链中暴露的预设错误，持续反制。",
        "四辩": "你扮演对方一辩反制，追问从价值层面挑战我方四辩的质询预设，试图守框架。",
    }

    guide = followup_guides.get(position, followup_guides["一辩"])

    prompt = f"""你正在模拟新国辩赛制中的质询环节。

辩题：{st.session_state.topic}
我方（你的对手）持方：{st.session_state.side}
我方立论框架：
{framework[:2000]}

以下是教练对上一轮接质的评价：
{evaluation}

根据评价决定下一步：
- 如果评价显示回答合格（6分以上），生成一个【追问】。{guide}
- 如果评价显示回答不合格（6分以下），生成一个【新问题】，换一个攻击方向。

注意：严格按照你扮演的对方对应辩位角色提问（四辩问前端、三辩问中端）。问题应简短有力。直接输出问题，不要任何其他文字。"""

    return llm.ask("你是对方辩手，你正在质询我方。根据教练评价决定追问还是换题。只输出问题。", prompt, temperature=0.85)


def update_framework_from_edit(position: str, edited_speech: str) -> str:
    """根据用户编辑后的辩词，更新立论框架并返回优化建议"""
    llm = get_llm()
    old_framework = st.session_state.arguments
    old_speech = st.session_state.position_results.get(position, "")

    prompt = f"""你是辩论教练。我方{position}对辩词进行了修改。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}

原始辩词：
{old_speech[:3000]}

修改后辩词：
{edited_speech[:3000]}

当前立论框架：
{old_framework[:3000]}

请完成以下任务：

1. **识别修改要点**：对比原始和修改后，提炼出用户新增/调整的核心论证内容
2. **更新立论框架**：如果修改涉及新论点、新定义或新判准，将这些补充到立论框架中，输出更新后的完整框架
3. **优化建议**：对修改后的辩词内容提出进一步优化建议

输出格式：
## 识别到的修改
（简要列出关键变化）

## 更新后的立论框架
（完整框架，保留原有未变内容 + 新增/调整内容）

## 优化建议
（对辩词的进一步打磨建议）"""

    result = llm.ask("你是辩论教练。审阅辩词修改，更新框架和建议。", prompt, temperature=0.7)

    # 尝试提取框架部分
    framework_marker = "## 更新后的立论框架"
    if framework_marker in result:
        new_framework = result.split(framework_marker)[1]
        next_marker = new_framework.find("\n## ")
        if next_marker > 0:
            new_framework = new_framework[:next_marker]
        st.session_state.arguments = new_framework.strip()

    return result


def _optimize_response(position: str, question: str, original: str, evaluation: str) -> str:
    """根据评价优化接质回应"""
    llm = get_llm()
    framework = st.session_state.arguments
    no_askback = "⚠️ 绝对不能反问对方，必须正面陈述。" if position in ("一辩", "二辩") else ""

    prompt = f"""你是辩论教练，请根据评价意见，优化我方{position}的接质回应。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
立论框架：
{framework[:2000]}

对方质询：{question}
原始回应：{original}

教练评价及改进建议：
{evaluation}

{no_askback}

请输出优化后的回应（40-100字）。必须：
1. 保持原回应的核心立场
2. 按评价意见修正不足
3. 更加精准、更有力地守住我方立场
4. {position}不能反问，只能正面陈述

直接输出优化后的回应文本，不要任何前缀。"""

    return llm.ask("你是辩论教练。优化接质回应，使其更精准有力。", prompt, temperature=0.7)


def _render_cross_exam_section():
    """渲染接质训练区域"""
    pos = st.session_state.current_position

    # 只有框架确认后才显示
    if not st.session_state.framework_reviewed:
        return

    st.markdown("---")
    st.markdown('<h2 class="section-header">🎙️ 接质训练</h2>', unsafe_allow_html=True)

    # 根据辩位显示不同说明
    pos_instructions = {
        "一辩": f"**{pos}前端接质训练**（一辩陈词后接对方四辩质询）\n\n"
                "🎭 **角色分配**：AI 扮演对方四辩向你提问 | 你是我方一辩，守护前端框架\n\n"
                "对方四辩攻击方向（前端问题）：\n"
                "- 定义自洽性：核心概念定义有无内部矛盾？\n"
                "- 判准公平性：判准是否偏向我方？能否覆盖辩题全部情况？\n"
                "- 前提漏洞：隐含前提能否举出反例？\n\n"
                "⚠️ 一辩是被质询方，不能反问。温和而坚定地守护框架根基。",
        "二辩": f"**{pos}中端接质训练**（申论驳论后接对方三辩质询）\n\n"
                "🎭 **角色分配**：AI 扮演对方三辩向你提问 | 你是我方二辩，守护中端战场\n\n"
                "对方三辩攻击方向（中端核心战场）：\n"
                "- 情景陷阱：构造具体案例逼你在该情景下让步\n"
                "- 例证可靠性：你方例子是否普遍？数据来源是否权威？\n"
                "- 逻辑延展：接受你的逻辑会导致什么荒谬推论？\n\n"
                "⚠️ 二辩是被质询方，不能反问。灵活化解陷阱，善于用例子反制例子。",
        "三辩": f"**{pos}反制应对训练**（质询对方中端战场时遭反制）\n\n"
                "🎭 **角色分配**：AI 扮演对方二辩反制你 | 你是我方三辩，正在质询对方中端战场\n\n"
                "训练目标：\n"
                "- 当你的问题链被对方反制时，如何冷静拆解对方的反制逻辑\n"
                "- 保持质询主动权，不被对方带偏\n"
                "- 在反制中继续推进中场收束\n\n"
                "⚠️ 三辩是主动质询方，但训练中模拟对方反制场景，锻炼应变能力。",
        "四辩": f"**{pos}反制应对训练**（质询对方前端问题时遭反制）\n\n"
                "🎭 **角色分配**：AI 扮演对方一辩反制你 | 你是我方四辩，正在质询对方前端问题\n\n"
                "训练目标：\n"
                "- 当对方一辩坚守框架时，如何绕过防线找到突破口\n"
                "- 将对方的反制转化为价值层面的对抗\n"
                "- 为后续的全局总结陈词积累素材\n\n"
                "⚠️ 四辩是主动质询方（质询对方前端），训练中模拟被反制场景，锻炼价值升维能力。",
    }
    st.info(pos_instructions.get(pos, pos_instructions["一辩"]))

    # 问题生成区
    col_q, col_r, col_h = st.columns([2, 1, 1])
    with col_q:
        btn_label = "🔄 刷新质询问题" if st.session_state.cross_exam_current_q else "🎯 生成质询问题"
        if st.button(btn_label, use_container_width=True):
            with st.spinner("正在生成质询问题..."):
                st.session_state.cross_exam_current_q = generate_cross_question(pos)
                st.session_state.cross_exam_need_help = False
            st.rerun()

    # 显示当前问题
    if st.session_state.cross_exam_current_q:
        st.markdown(f"""
        <div style="background:rgba(201,169,110,0.06); color:#E8E6F0; border:1px solid rgba(201,169,110,0.2); border-radius:10px; padding:16px; margin:12px 0;">
            <strong>🎯 对方辩手向你提问：</strong><br>
            <span style="font-size:1.2em;">{st.session_state.cross_exam_current_q}</span>
        </div>
        """, unsafe_allow_html=True)

        # 用户回答区
        user_answer = st.text_area(
            "✍️ 你的回应：",
            placeholder="输入你的回应...（不能反问，必须正面回答）",
            key="cross_answer_input",
            height=80,
        )

        col_a1, col_a2, col_a3 = st.columns(3)
        with col_a1:
            if st.button("📤 提交回应", use_container_width=True, disabled=not user_answer):
                with st.spinner("教练正在评价你的回应..."):
                    evaluation = evaluate_cross_answer(pos, st.session_state.cross_exam_current_q, user_answer)
                # 检测论据缺失
                gap = detect_evidence_gap(f"质询问题：{st.session_state.cross_exam_current_q}\n回应：{user_answer}")
                if gap:
                    auto_fill_evidence(gap.replace("缺少论据：", "").strip())
                st.session_state.cross_exam_rounds.append({
                    "position": pos,
                    "question": st.session_state.cross_exam_current_q,
                    "user_answer": user_answer,
                    "evaluation": evaluation,
                })
                st.session_state.cross_exam_evaluation = evaluation
                st.session_state.cross_exam_last_answer = user_answer
                st.session_state.cross_exam_current_q = ""
                st.rerun()
        with col_a2:
            if st.button("🤖 AI 帮我回应（学习）", use_container_width=True):
                with st.spinner("AI 正在生成示范回应..."):
                    ai_answer = generate_cross_response(pos, st.session_state.cross_exam_current_q)
                    evaluation = evaluate_cross_answer(pos, st.session_state.cross_exam_current_q, ai_answer)
                st.session_state.cross_exam_rounds.append({
                    "position": pos,
                    "question": st.session_state.cross_exam_current_q,
                    "ai_answer": ai_answer,
                    "evaluation": evaluation,
                })
                st.session_state.cross_exam_evaluation = "**【AI示范回应】**\n" + evaluation
                st.session_state.cross_exam_current_q = ""
                st.rerun()
        with col_a3:
            if st.button("⏭️ 跳过", use_container_width=True):
                st.session_state.cross_exam_rounds.append({
                    "position": pos,
                    "question": st.session_state.cross_exam_current_q,
                    "user_answer": "[跳过]",
                    "evaluation": "",
                })
                st.session_state.cross_exam_current_q = ""
                st.rerun()

    # 评价结果 + 下一步按钮
    if st.session_state.cross_exam_evaluation:
        st.markdown(f"""
        <div style="background:rgba(108,92,231,0.05); color:#E8E6F0; border:1px solid rgba(108,92,231,0.12); border-radius:10px; padding:16px; margin:12px 0;">
            <strong>📊 教练评价：</strong><br>
            {st.session_state.cross_exam_evaluation}
        </div>
        """, unsafe_allow_html=True)

        if pos in ("一辩", "二辩"):
            st.caption("⚠️ 提醒：一辩/二辩是被质询方，没有追问对方的权力。下方按钮是让AI对手提出下一问，不是让你追问。")

        col_opt, col_next, col_clear = st.columns([1, 1.5, 0.8])
        with col_opt:
            if st.button("✨ 基于评价优化我的回应", use_container_width=True):
                with st.spinner("AI 正在优化回应..."):
                    optimized = _optimize_response(
                        pos, st.session_state.cross_exam_rounds[-1]["question"],
                        st.session_state.cross_exam_rounds[-1].get("user_answer") or
                        st.session_state.cross_exam_rounds[-1].get("ai_answer", ""),
                        st.session_state.cross_exam_evaluation)
                    st.session_state.cross_exam_rounds[-1]["optimized_answer"] = optimized
                    st.session_state.cross_exam_evaluation += f"\n\n**✨ 优化后回应**：{optimized}"
                st.rerun()
        with col_next:
            if st.button("▶️ 下一步（AI 对手决定追问或换题）", type="primary", use_container_width=True):
                with st.spinner("AI 正在决定下一步..."):
                    next_q = decide_next_question(pos, st.session_state.cross_exam_evaluation)
                    st.session_state.cross_exam_current_q = next_q
                    st.session_state.cross_exam_evaluation = ""
                    st.session_state.cross_exam_last_answer = ""
                st.rerun()
        with col_clear:
            if st.button("🗑️ 结束本轮训练", use_container_width=True):
                st.session_state.cross_exam_evaluation = ""
                st.rerun()

    # 历史记录
    if st.session_state.cross_exam_rounds:
        with st.expander(f"📜 质询记录（{len(st.session_state.cross_exam_rounds)} 轮）", expanded=False):
            for i, r in enumerate(st.session_state.cross_exam_rounds):
                st.markdown(f"**第 {i+1} 轮**")
                st.markdown(f"🎯 对方问：{r['question']}")
                if r.get("user_answer") and r["user_answer"] != "[跳过]":
                    st.markdown(f"🧑 你答：{r['user_answer']}")
                if r.get("ai_answer"):
                    st.markdown(f"🤖 AI示范答：{r['ai_answer']}")
                if r.get("optimized_answer"):
                    st.markdown(f"✨ 优化后：{r['optimized_answer']}")
                if r.get("evaluation"):
                    st.markdown(f"📊 评价：{r['evaluation'][:300]}")
                st.markdown("---")

            if st.button("🗑️ 清空质询记录", use_container_width=True):
                st.session_state.cross_exam_rounds = []
                st.session_state.cross_exam_current_q = ""
                st.session_state.cross_exam_evaluation = ""
                st.session_state.editing_speech = False
                st.session_state.speech_feedback = ""
                st.session_state.interrogate_q = ""
                st.session_state.interrogate_ai_a = ""
                st.session_state.interrogate_eval = ""
                st.rerun()


# ============================================================
# 三辩/四辩 质询训练（主动提问方）
# ============================================================
def generate_question_to_ask(position: str, direction: str = "") -> str:
    """为三辩/四辩生成一个建议的质询问题，可指定方向"""
    llm = get_llm()
    framework = st.session_state.arguments
    history_context = ""
    if st.session_state.interrogate_history:
        recent = st.session_state.interrogate_history[-3:]
        history_context = "此前质询记录：\n" + "\n".join(
            f"Q: {r['q']}\nA: {r['a']}" for r in recent
        )

    guides = {
        "三辩": """你代表我方立场，基于我方立论框架设计质询问题，质询对方中端战场。

核心原则：问题必须从我方框架出发——用我方的定义、判准、论点作为武器，攻击对方中端论证的薄弱之处。

具体策略：
1. 基于我方核心论点，设计封闭式问题逼对方在我方有利的框架下回答
2. 用我方框架中的逻辑链去检验对方的例证——对方的例子在我方逻辑下是否成立？
3. 基于我方判准，质询对方的论证是否满足同一标准
4. 归谬/类比必须以我方立场为出发点
问题风格：简短封闭式（15-25字），让我方框架的优势在问题中自然显现。""",
        "四辩": """你代表我方立场，基于我方立论框架设计质询问题，质询对方前端问题（定义/判准/前提）。

核心原则：问题必须以我方框架为根基——用我方的定义去挑战对方的定义，用我方的判准去评估对方的判准。

四辩核心技巧——陷阱式共识偷渡：
这是四辩最高级的质询技巧。设计看似中立甚至无害的问题，让对方在回答"是"的过程中，不知不觉接受了我方框架的前提。具体方法：
- 先问一个对方无法否认的常识性前提 → 对方同意
- 再问一个该前提在我方框架下的自然延伸 → 对方难以拒绝
- 对方在连续同意后，已被拉入我方框架
- 此时再问关键问题，对方已失去退路
示例：不问"你是否同意AI效率更高？"（对方会警惕），而是问"效率是不是决策的一个参考因素？"（无法否认）→"既然效率是一个因素，那在效率差距显著时，是否应该优先考虑效率？"（已在你的框架内）

具体策略：
1. 陷阱共识偷渡：设计2-3个递进的封闭式问题，让对方在不知不觉中接受我方前提
2. 基于我方定义，质询对方定义是否能通过我方定义的合理标准
3. 基于我方判准，挑战对方判准的公平性和覆盖面
4. 逼对方在"我方框架 vs 对方框架"的对比中暴露其框架的缺陷
问题风格：看似中立实则设伏（20-30字），让对方回答"是"的瞬间就进入我方框架。""",
    }
    guide = guides.get(position, guides["三辩"])

    direction_instruction = ""
    if direction:
        direction_instruction = f"""
【指定质询方向】
请专门按照以下方向设计问题：{direction}
忽略通用策略，聚焦于此方向。"""

    prompt = f"""你是我方{position}，正在基于我方立论框架准备质询对方。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}

【我方立论框架（你的问题必须根植于此）】
{framework[:3000]}

{history_context}

【你的质询策略】
{guide}
{direction_instruction}

重要：你提出的问题必须体现我方框架的核心主张。不是在真空中提问，而是用我方的定义、判准、论点作为问题的根基。
请生成一个质询问题。直接输出问题，不要其他文字。"""

    return llm.ask("你是我方辩手，基于我方框架设计质询问题。问题必须体现我方立场和论证优势。", prompt, temperature=0.9)


def simulate_opponent_answer(position: str, question: str) -> str:
    """AI模拟对方辩手对质询问题的回应"""
    llm = get_llm()
    framework = st.session_state.arguments

    opponent_roles = {
        "三辩": "你是对方二辩（申论驳论手）。我方三辩正在质询你的中端论证和例证。你必须正面回应，不能反问。根据你的立场给出回应。",
        "四辩": "你是对方一辩（立论陈词手）。我方四辩正在质询你的前端定义和判准。你必须坚守你的框架，正面回应，不能反问。",
    }

    prompt = f"""你正在模拟新国辩赛制中被质询的对方辩手。

辩题：{st.session_state.topic}
对方（我方的对手）持方：{'反方' if st.session_state.side == '正方' else '正方'}
对方立论框架（你需要从对方角度防守）：
{framework[:2000]}

我方质询问题：{question}

{opponent_roles.get(position, opponent_roles['三辩'])}

请给出对方的回应（30-80字）。必须正面回应，不能反问。直接输出回应文本。"""

    return llm.ask("你是被质询的对方辩手。正面回应质询问题。", prompt, temperature=0.8)


def evaluate_question_quality(position: str, question: str, response: str) -> str:
    """评价质询问题的质量"""
    llm = get_llm()
    framework = st.session_state.arguments

    criteria = {
        "三辩": "- 框架根植性：问题是否源自我方立论框架（而非天外飞仙式提问）？\n- 攻击精准度：是否打在对方中端论证的薄弱点？\n- 封闭性：是否控制了对方回答空间？",
        "四辩": "- 框架根植性：问题是否用我方定义/判准去审视对方？\n- 陷阱设计力：问题是否有共识偷渡效果？对方是否难以拒绝？\n- 递进空间：该问题是否为后续追问留了递进拉入的空间？",
    }

    prompt = f"""你是资深辩论教练，正在评价我方{position}的质询问题。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
立论框架：
{framework[:2000]}

我方质询问题：{question}
对方回应：{response}

评价维度：
{criteria.get(position, criteria['三辩'])}

请给出：
**评分**（满分10分）：X/10
**优点**：
**可改进**：
**追问建议**（基于对方的回应，下一步追问什么？）：

简洁输出，150字以内。直接给出评价。"""

    return llm.ask("你是辩论教练。评价质询问题的质量，给出追问建议。", prompt, temperature=0.7)


def _render_interrogate_section():
    """渲染三辩/四辩的质询训练"""
    pos = st.session_state.current_position

    if not st.session_state.framework_reviewed:
        return

    st.markdown("---")
    st.markdown('<h2 class="section-header">🎙️ 质询训练（主动提问）</h2>', unsafe_allow_html=True)

    if pos == "三辩":
        st.info(f"**{pos}质询训练**：你是主动质询方，质询对方中端核心战场。AI 扮演对方二辩回应你的质询。\n\n"
                "训练目标：设计精准的问题链，攻击对方申论论证和例证可靠性，为后续中场收束打基础。")
    else:
        st.info(f"**{pos}质询训练**：你是主动质询方，质询对方前端问题。AI 扮演对方一辩回应你的质询。\n\n"
                "训练目标：攻击对方定义/判准/前提，引导评委看到对方框架的根本缺陷，为后续全局总结铺垫。")

    # 方向设定
    direction_presets = {
        "三辩": ["自动（AI 自行选择）", "归谬法：假设对方成立推导荒谬结论",
                 "类比法：用类似情景揭示逻辑问题", "挑战例证：攻击对方数据的代表性和来源",
                 "逻辑断层：找论证链条中的薄弱环节", "特例法：找对方论证无法解释的特例"],
        "四辩": ["自动（AI 自行选择）", "陷阱共识偷渡：先让对方同意一个看似无害的前提，逐步拉入我方框架",
                 "攻定义自洽性：对方核心概念有无内部矛盾",
                 "攻判准公平性：对方判准是否偏袒其立场",
                 "挖前提反例：对方立论前提能否找到反例",
                 "逼定义收缩：用封闭式问题逼对方让步",
                 "价值审视：对方框架的价值预设是否合理"],
    }
    presets = direction_presets.get(pos, direction_presets["三辩"])
    use_custom_dir = st.checkbox("🎯 设定质询方向", key="set_direction")
    direction = ""
    if use_custom_dir:
        col_preset, col_custom = st.columns([1, 1])
        with col_preset:
            preset = st.selectbox("预设方向", presets, key="dir_preset")
            if preset != "自动（AI 自行选择）":
                direction = preset
        with col_custom:
            custom_dir = st.text_input("或自定义方向", placeholder="例如：专攻对方关于成本效益的论证",
                                       key="dir_custom")
            if custom_dir:
                direction = custom_dir

    # 问题输入区
    col_gen, col_manual = st.columns([1, 2])
    with col_gen:
        btn_label = f"🤖 AI 生成问题" + (f"（{direction[:15]}...）" if direction else "")
        if st.button(btn_label, use_container_width=True):
            with st.spinner("正在生成质询问题..."):
                st.session_state.interrogate_q = generate_question_to_ask(pos, direction)
                st.session_state.interrogate_ai_a = ""
                st.session_state.interrogate_eval = ""
            st.rerun()
    with col_manual:
        user_q = st.text_input("✍️ 或自己输入质询问题：", placeholder="输入你想问对方的问题...",
                               key="interrogate_input")
        if st.button("📤 提交问题", disabled=not user_q):
            st.session_state.interrogate_q = user_q
            st.session_state.interrogate_ai_a = ""
            st.session_state.interrogate_eval = ""
            st.rerun()

    # 显示问题和AI回应
    if st.session_state.interrogate_q:
        st.markdown(f"""
        <div style="background:rgba(162,155,254,0.05); color:#E8E6F0; border:1px solid rgba(162,155,254,0.12); border-radius:10px; padding:16px; margin:12px 0;">
            <strong>🎯 你的质询问题：</strong><br>
            <span style="font-size:1.15em;">{st.session_state.interrogate_q}</span>
        </div>
        """, unsafe_allow_html=True)

        if not st.session_state.interrogate_ai_a:
            if st.button("▶️ 让 AI（对方）回应", type="primary", use_container_width=True):
                with st.spinner("AI 正在模拟对方回应..."):
                    st.session_state.interrogate_ai_a = simulate_opponent_answer(pos, st.session_state.interrogate_q)
                st.rerun()
        else:
            st.markdown(f"""
            <div style="background:rgba(201,169,110,0.06); color:#E8E6F0; border:1px solid rgba(201,169,110,0.2); border-radius:10px; padding:16px; margin:12px 0;">
                <strong>🗣️ 对方回应：</strong><br>
                {st.session_state.interrogate_ai_a}
            </div>
            """, unsafe_allow_html=True)

            # 评价（可选）
            if st.session_state.interrogate_eval:
                st.markdown(f"""
                <div style="background:rgba(108,92,231,0.05); color:#E8E6F0; border:1px solid rgba(108,92,231,0.12); border-radius:10px; padding:16px; margin:12px 0;">
                    <strong>📊 教练评价：</strong><br>
                    {st.session_state.interrogate_eval}
                </div>
                """, unsafe_allow_html=True)

            # 操作按钮（全部平铺，不绑定评价依赖）
            col_eval, col_opt, col_fw_self, col_fw_ai = st.columns(4)
            with col_eval:
                if st.button("📊 评价问题" if not st.session_state.interrogate_eval else "📊 重新评价",
                             use_container_width=True):
                    with st.spinner("教练正在评价..."):
                        st.session_state.interrogate_eval = evaluate_question_quality(
                            pos, st.session_state.interrogate_q, st.session_state.interrogate_ai_a)
                    st.rerun()
            with col_opt:
                if st.button("✨ 优化问题", use_container_width=True,
                            disabled=not st.session_state.interrogate_eval,
                            help="需要先评价才能优化"):
                    with st.spinner("AI 正在优化问题..."):
                        llm = get_llm()
                        framework = st.session_state.arguments
                        optimized_q = llm.ask(
                            "你是我方辩手。根据教练评价优化质询问题。优化后的问题必须更精准、更根植于我方的立论框架。只输出优化后的问题。",
                            f"辩题：{st.session_state.topic}\n持方：{st.session_state.side}\n我方框架：{framework[:1500]}\n原始问题：{st.session_state.interrogate_q}\n教练评价：{st.session_state.interrogate_eval}\n\n请输出优化后的问题：",
                            temperature=0.7)
                        st.session_state.interrogate_q = optimized_q.strip()
                        st.session_state.interrogate_ai_a = ""
                        st.session_state.interrogate_eval = ""
                    st.rerun()
            with col_fw_self:
                if st.button("✍️ 追问（自输入）", use_container_width=True):
                    st.session_state.interrogate_history.append({
                        "q": st.session_state.interrogate_q,
                        "a": st.session_state.interrogate_ai_a,
                    })
                    st.session_state.interrogate_q = ""
                    st.session_state.interrogate_ai_a = ""
                    st.session_state.interrogate_eval = ""
                    st.rerun()
            with col_fw_ai:
                if st.button("🤖 AI生成追问", use_container_width=True):
                    st.session_state.interrogate_history.append({
                        "q": st.session_state.interrogate_q,
                        "a": st.session_state.interrogate_ai_a,
                    })
                    with st.spinner("生成追问..."):
                        llm = get_llm()
                        framework = st.session_state.arguments
                        fw = llm.ask(
                            "你是我方辩手。基于我方框架和上一轮质询，生成追问。追问必须从我方立场出发。只输出问题。",
                            f"辩题：{st.session_state.topic}\n持方：{st.session_state.side}\n我方框架：{framework[:1500]}\n上一问：{st.session_state.interrogate_q}\n对方回应：{st.session_state.interrogate_ai_a}\n\n基于我方框架和对方回应生成追问：",
                            temperature=0.8)
                        st.session_state.interrogate_q = fw.strip()
                        st.session_state.interrogate_ai_a = ""
                        st.session_state.interrogate_eval = ""
                    st.rerun()

            col_new, col_rec, _, _ = st.columns(4)
            with col_new:
                if st.button("🆕 新问题", use_container_width=True):
                    st.session_state.interrogate_history.append({
                        "q": st.session_state.interrogate_q,
                        "a": st.session_state.interrogate_ai_a,
                    })
                    st.session_state.interrogate_q = ""
                    st.session_state.interrogate_ai_a = ""
                    st.session_state.interrogate_eval = ""
                    st.rerun()
            with col_rec:
                if st.button("📝 记录", use_container_width=True):
                    st.session_state.interrogate_history.append({
                        "q": st.session_state.interrogate_q,
                        "a": st.session_state.interrogate_ai_a,
                    })
                    st.session_state.interrogate_q = ""
                    st.session_state.interrogate_ai_a = ""
                    st.session_state.interrogate_eval = ""
                    st.rerun()

    # 历史
    if st.session_state.interrogate_history:
        with st.expander(f"📜 质询记录（{len(st.session_state.interrogate_history)} 轮）", expanded=False):
            for i, r in enumerate(st.session_state.interrogate_history):
                st.markdown(f"**第 {i+1} 轮**")
                st.markdown(f"🎯 你问：{r['q']}")
                st.markdown(f"🗣️ 对方答：{r['a']}")
                st.markdown("---")
            if st.button("🗑️ 清空质询记录", key="clear_interrogate"):
                st.session_state.interrogate_history = []
                st.rerun()


# ============================================================
# 论据库模块
# ============================================================
def extract_evidence_from_text(text: str, source: str) -> list[dict]:
    """从文本中提取论据条目"""
    llm = get_llm()
    prompt = f"""你是辩论资料研究员。请从以下文本中提取可用于辩论的论据条目。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}

文本内容：
{text[:8000]}

请提取其中的关键论据，每条论据包含：
- 具体事实/数据/案例/名言
- 每个条目 40-200 字
- 标注类型（事实数据/案例/权威引述/逻辑论证）

输出 JSON 数组：[{{"content": "...", "type": "..."}}]
如果没有可用论据，输出空数组 []。直接输出 JSON。"""

    result = llm.ask("你是辩论资料研究员。提取论据，输出 JSON。", prompt, temperature=0.6)
    try:
        start = result.find("[")
        end = result.rfind("]") + 1
        if start >= 0 and end > start:
            items = json.loads(result[start:end])
            import datetime
            now = datetime.datetime.now().strftime("%m-%d %H:%M")
            for item in items:
                item["source"] = source
                item["added_at"] = now
                item["tags"] = [item.get("type", "论据")]
            return items
    except Exception:
        pass
    # 解析失败，将整段作为一条论据
    import datetime
    return [{"content": text[:500], "source": source, "tags": ["论据"],
             "added_at": datetime.datetime.now().strftime("%m-%d %H:%M")}]


def add_evidence_manual(content: str, source: str = "手动添加"):
    """手动添加论据"""
    import datetime
    st.session_state.evidence_library.append({
        "content": content.strip(),
        "source": source,
        "tags": ["手动添加"],
        "added_at": datetime.datetime.now().strftime("%m-%d %H:%M"),
    })


def search_web_for_evidence(claim: str) -> str:
    """搜索网页补充论据"""
    if has_exa():
        results = _exa_search(f"{st.session_state.topic} {claim} 数据 研究 案例", 3)
        texts = [r["text"][:800] for r in results if r.get("text")]
        return "\n".join(texts) if texts else ""
    return ""


def detect_evidence_gap(text: str) -> str | None:
    """检测文本中是否缺少论据支撑，返回缺少论据的主张"""
    llm = get_llm()
    prompt = f"""你是辩论论据审查员。检查以下文本是否有充分的论据支撑。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}

文本：
{text[:2000]}

检查要点：
- 每个事实性主张是否有数据/案例支撑
- 每个因果推论是否有实证依据
- 是否有"裸主张"（只有结论没有论据）

如果所有主张都有充分论据，回复"无缺失"。
如果有缺失，回复："缺少论据：<具体主张>"（指出最需要补充论据的那个主张）。

简洁输出，直接给结论。"""

    result = llm.ask("你是论据审查员。简洁判断论据是否充分。", prompt, temperature=0.5)
    if "无缺失" in result or "充分" in result:
        return None
    return result.strip()


def auto_fill_evidence(claim: str) -> bool:
    """自动搜索并补充论据，返回是否成功添加"""
    with st.spinner(f"🔍 正在搜索论据：「{claim[:50]}」..."):
        web_text = search_web_for_evidence(claim)
        if web_text:
            items = extract_evidence_from_text(web_text, f"网页搜索：{claim[:50]}")
            added = 0
            for item in items:
                # 去重
                existing = [e["content"][:80] for e in st.session_state.evidence_library]
                if item["content"][:80] not in existing:
                    st.session_state.evidence_library.append(item)
                    added += 1
            if added > 0:
                st.session_state.evidence_notification = (
                    f"🔔 论据库已自动更新：为「{claim[:40]}」补充了 {added} 条论据。"
                    f"来源：{items[0]['source'] if items else '网页搜索'}"
                )
                return True
        return False


def _render_evidence_panel():
    """渲染论据库面板"""
    st.markdown("---")
    st.markdown('<h2 class="section-header">📚 论据库</h2>', unsafe_allow_html=True)

    # 通知
    if st.session_state.evidence_notification:
        st.toast(st.session_state.evidence_notification, icon="🔔")
        st.info(st.session_state.evidence_notification)
        if st.button("标记已读", key="clear_notification"):
            st.session_state.evidence_notification = ""
            st.rerun()

    lib = st.session_state.evidence_library
    st.caption(f"共 {len(lib)} 条论据 | 可用于质询和接质训练")

    # 添加论据
    tab_add, tab_browse = st.tabs(["➕ 添加论据", f"📋 浏览论据（{len(lib)}）"])

    with tab_add:
        # 上传文件
        uf = st.file_uploader("上传资料提取论据", type=["pdf", "docx", "txt", "md"],
                              key="evidence_upload", help="上传 PDF/Word/文本，AI 自动提取论据")
        if uf and st.button("📥 提取论据", key="extract_evidence_btn"):
            with st.spinner("正在提取论据..."):
                material = read_uploaded(uf)
                items = extract_evidence_from_text(material, f"文件：{uf.name}")
                for item in items:
                    st.session_state.evidence_library.append(item)
                st.success(f"从 {uf.name} 提取了 {len(items)} 条论据！")
                st.rerun()

        # 手动输入
        st.markdown("**或手动输入：**")
        manual = st.text_area("输入论据内容", placeholder="粘贴一段数据、案例或权威引述...",
                              key="manual_evidence", height=80)
        col_add, col_search = st.columns(2)
        with col_add:
            if st.button("➕ 添加", use_container_width=True, disabled=not manual):
                add_evidence_manual(manual)
                st.success("论据已添加！")
                st.rerun()
        with col_search:
            if st.button("🔍 搜索网页补充", use_container_width=True,
                         disabled=not manual or not has_exa()):
                items = extract_evidence_from_text(manual, "手动输入+网页搜索")
                for item in items:
                    st.session_state.evidence_library.append(item)
                added = len(items)
                # 同时搜索网页
                web_text = search_web_for_evidence(manual[:100])
                if web_text:
                    web_items = extract_evidence_from_text(web_text, f"网页搜索：{manual[:50]}")
                    for item in web_items:
                        st.session_state.evidence_library.append(item)
                    added += len(web_items)
                st.success(f"添加了 {added} 条论据！")
                st.rerun()

    with tab_browse:
        if not lib:
            st.info("论据库为空。上传资料或手动添加论据。")
        else:
            # 搜索过滤
            search = st.text_input("🔍 搜索论据", key="search_evidence",
                                   placeholder="输入关键词过滤...")
            filtered = lib
            if search:
                filtered = [e for e in lib if search in e.get("content", "")]
                st.caption(f"找到 {len(filtered)} 条匹配")

            for i, ev in enumerate(reversed(filtered[-20:])):  # 最近20条
                real_idx = len(lib) - 1 - lib[::-1].index(ev) if ev in lib else i
                cols = st.columns([0.05, 0.75, 0.1, 0.1])
                with cols[0]:
                    tag_emoji = {"事实数据": "📊", "案例": "📖", "权威引述": "🎓", "逻辑论证": "🔗",
                                 "手动添加": "✏️", "论据": "📌"}
                    emoji = tag_emoji.get(ev.get("tags", ["论据"])[0], "📌")
                    st.markdown(emoji)
                with cols[1]:
                    st.markdown(f"{ev['content'][:200]}")
                    st.caption(f"来源：{ev.get('source','')} | {ev.get('added_at','')}")
                with cols[2]:
                    # 复制按钮用 markdown
                    pass
                with cols[3]:
                    if st.button("🗑️", key=f"del_ev_{real_idx}", help="删除"):
                        st.session_state.evidence_library.pop(real_idx)
                        st.rerun()

            if len(lib) > 20:
                st.caption(f"... 显示最近 20 条，共 {len(lib)} 条。使用搜索过滤。")

            if st.button("🗑️ 清空论据库", use_container_width=True):
                st.session_state.evidence_library = []
                st.rerun()

    # 论据转辩词
    if lib:
        st.markdown("---")
        st.markdown("### 🎤 论据转化辩词")
        col_btn, col_target = st.columns([1, 1])
        with col_target:
            target = st.selectbox("转化目标", ["通用攻防用语", "一辩立论引用", "二辩对辩反击",
                                              "三辩质询素材", "四辩总结升华"],
                                  key="evidence_target")
        with col_btn:
            if st.button("🤖 将论据转化为场上辩词", type="primary", use_container_width=True):
                with st.spinner("正在将论据转化为可用的辩词..."):
                    result = _convert_evidence_to_speech(target)
                    st.session_state._last_converted_speech = result
                st.rerun()

        if st.session_state.get("_last_converted_speech"):
            st.markdown(st.session_state._last_converted_speech)
            st.download_button("📥 导出转化辩词",
                               data=generate_docx("论据转化辩词", st.session_state._last_converted_speech),
                               file_name=f"论据转化_{target}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)


def _convert_evidence_to_speech(target: str) -> str:
    """将论据库中的论据转化为可供赛场使用的辩词，保留学理逻辑和来源引用"""
    llm = get_llm()
    framework = st.session_state.arguments[:3000]

    # 完整传递论据，不截断关键内容
    evidence_items = []
    for e in st.session_state.evidence_library[-20:]:
        tags_str = "·".join(e.get('tags', ['论据']))
        src = e.get('source', '未知来源')
        evidence_items.append(f"[{tags_str}] 来源：{src}\n内容：{e['content'][:500]}")
    evidence_text = "\n\n---\n\n".join(evidence_items)

    prompt = f"""你是顶级辩论教练，擅长将学术论据转化为有深度、有学理的场上辩词。

辩题：{st.session_state.topic}
我方持方：{st.session_state.side}
转化目标：{target}

我方立论框架：
{framework}

论据素材（含来源）：
{evidence_text[:8000]}

转化要求（极其重要）：
1. **保留学理深度**：每条论据背后的学术原理、哲学逻辑、因果机制必须体现在辩词中，不能只摘表面结论。例如引用一个研究数据时，要解释这个数据为什么能推出我方结论（讲清"所以然"）
2. **明确引用来源**：在辩词中口头化地融入来源信息。例如："哈佛大学2023年的一项追踪研究显示..."、"正如罗尔斯在《正义论》中指出的..."
3. **构建完整逻辑链**：每个辩词片段必须是"主张 → 论据引用 → 学理解释 → 推论 → 扣回立场"的完整闭环
4. **语言口语化但不失深度**：是赛场上能说的自然口语，但内容要有学术重量，让评委感受到论据的扎实程度
5. **区分场景适配**：根据目标场景调整语气和长度

输出格式：
### 📐 学理框架总览
（2-3句话概括：这批论据背后的核心学术逻辑是什么，在辩论中如何运用）

**辩词片段 1**（适用于<具体场景>）
> <完整辩词正文，包含引用+学理+推理>
📎 引用：<具体来源>

**辩词片段 2**（适用于<具体场景>）
> <完整辩词正文>
📎 引用：<具体来源>

...

### 🔗 论据间逻辑关联
（简要说明这些论据如何串联成一个完整的论证体系）

直接输出。"""

    return llm.ask("你是顶级辩论教练。将论据转化为有学理深度、有明确引用的场上辩词。保留论据背后的学术逻辑，不丢弃来源信息。", prompt, temperature=0.8)


# ---- 页面3：分辩位训练 ----
def render_train_page():
    pos = st.session_state.current_position
    profile = POSITION_PROFILES.get(pos, POSITION_PROFILES["一辩"])

    st.markdown(f'<h1 class="main-title">{pos}训练</h1>', unsafe_allow_html=True)
    st.markdown(f'<p class="sub-title">{st.session_state.topic} · {st.session_state.side}方 · {profile["role"]}</p>',
                unsafe_allow_html=True)

    # 立论框架参考（始终可见）
    if st.session_state.standpoint_confirmed:
        with st.expander("📋 查看立论框架（随时参考）", expanded=False):
            st.markdown(st.session_state.arguments)
        if st.session_state.framework_reviewed and st.session_state.position_tasks.get("full"):
            with st.expander(f"🎯 {pos}辩位任务", expanded=False):
                st.markdown(st.session_state.position_tasks["full"])

    with st.expander(f"📖 {pos}职责说明", expanded=False):
        st.markdown(f"**角色：** {profile['role']}")
        st.markdown(f"**任务：** {profile['duty']}")
        st.markdown(f"**风格：** {profile['style']}")

    c1, c2, c3, c4 = st.columns(4)
    for i, (p, role, _) in enumerate([
        ("一辩", "立论+前端接质", "框架搭建"),
        ("二辩", "申论驳论+中端接质", "攻防核心"),
        ("三辩", "质询+中场收束", "战场总结"),
        ("四辩", "全局质询+价值收束", "定音之锤"),
    ]):
        with [c1, c2, c3, c4][i]:
            if p == pos:
                st.success(f"✅ 当前：{p}（{role}）")
            else:
                if st.button(f"切换到 {p}", key=f"sw_{p}", use_container_width=True):
                    st.session_state.current_position = p
                    st.session_state.chat_history = []
                    st.session_state.cross_exam_rounds = []
                    st.session_state.cross_exam_current_q = ""
                    st.session_state.cross_exam_evaluation = ""
                    st.rerun()

    st.markdown("---")

    # 新国辩初稿生成
    if st.session_state.framework_reviewed and pos not in st.session_state.position_results:
        if st.button(f"🚀 基于立论框架生成{pos}初稿辩词（新国辩赛制）", type="primary", use_container_width=True):
            with st.spinner(f"正在为{pos}生成初稿辩词..."):
                speech = generate_position_speech(pos)
                st.session_state.position_speeches[pos] = speech
                st.session_state.position_results[pos] = speech
                st.session_state.chat_history = [
                    {"role": "user", "content": f"[基于立论框架生成{pos}初稿]"},
                    {"role": "assistant", "content": speech},
                ]
            save_session()
            st.rerun()

    # 训练模式
    mode = st.radio("选择训练模式：",
                    ["📄 上传资料（PDF/Word/文本）", "✍️ 输入想法进行优化"],
                    horizontal=True)
    trainer = PositionTrainer()

    if "上传资料" in mode:
        uf = st.file_uploader("上传你的资料文件", type=["pdf", "docx", "txt", "md"],
                              help="支持 PDF、Word、文本文件")
        if uf and st.button("📥 读取资料并生成辩词", type="primary", use_container_width=True):
            with st.spinner(f"正在读取资料并生成{pos}辩词..."):
                material = read_uploaded(uf)
                result = trainer.train_from_material(
                    pos, st.session_state.topic, st.session_state.side,
                    st.session_state.plan, material)
                st.session_state.position_results[pos] = result
                st.session_state.chat_history = [
                    {"role": "user", "content": f"[上传了资料：{uf.name}]"},
                    {"role": "assistant", "content": result},
                ]
                save_session()
            st.rerun()
    else:
        _render_chat(trainer)

    if pos in st.session_state.position_results:
        result = st.session_state.position_results[pos]
        st.markdown("---")
        st.markdown('<h2 class="section-header">辩词结果</h2>', unsafe_allow_html=True)

        if st.session_state.get("editing_speech"):
            # 编辑模式
            edited = st.text_area("✏️ 自由编辑辩词", value=result, height=500,
                                  key=f"edit_speech_{pos}")
            col_save, col_cancel = st.columns(2)
            with col_save:
                if st.button("💾 保存并更新框架", type="primary", use_container_width=True):
                    with st.spinner("正在分析修改并更新框架..."):
                        feedback = update_framework_from_edit(pos, edited)
                        st.session_state.position_results[pos] = edited
                        st.session_state.speech_feedback = feedback
                        st.session_state.editing_speech = False
                    save_session()
                    st.rerun()
            with col_cancel:
                if st.button("取消编辑", use_container_width=True):
                    st.session_state.editing_speech = False
                    st.rerun()

            # 显示反馈
            if st.session_state.speech_feedback:
                with st.expander("📊 框架更新与优化建议", expanded=True):
                    st.markdown(st.session_state.speech_feedback)
                    if st.button("🗑️ 关闭反馈"):
                        st.session_state.speech_feedback = ""
                        st.rerun()
        else:
            # 查看模式
            st.markdown(f'<div class="plan-box">{result}</div>', unsafe_allow_html=True)

            col_edit, col_dl = st.columns(2)
            with col_edit:
                if st.button("✏️ 编辑辩词", use_container_width=True):
                    st.session_state.editing_speech = True
                    st.session_state.speech_feedback = ""
                    st.rerun()
            with col_dl:
                st.download_button("📥 导出为 Word 文档",
                                   data=generate_docx(f"{pos}辩词", result),
                                   file_name=f"{pos}辩词_{st.session_state.topic}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True)

    # 训练模块：一辩/二辩接质训练，三辩/四辩质询训练
    if pos in ("一辩", "二辩"):
        _render_cross_exam_section()
    else:
        _render_interrogate_section()


def _render_chat(trainer: PositionTrainer):
    pos = st.session_state.current_position
    for msg in st.session_state.chat_history:
        role = msg["role"]
        cls = "chat-bubble-user" if role == "user" else "chat-bubble-ai"
        label = "🧑 你" if role == "user" else f"🤖 {pos}教练"
        st.markdown(f"""<div class="{cls}"><strong>{label}</strong><br>{msg['content'][:2000]}</div>""",
                    unsafe_allow_html=True)

    st.markdown("---")
    ui = st.text_area("✍️ 输入你的想法、草稿或论点片段",
                      placeholder="例如：我认为应该从效率的角度论证...\n或者直接粘贴你写的辩词草稿...",
                      height=150, key=f"user_input_{pos}")

    if st.button("✨ 优化为辩词", type="primary", use_container_width=True, disabled=not ui):
        with st.spinner(f"教练正在优化你的{pos}辩词..."):
            result = trainer.refine_input(
                pos, st.session_state.topic, st.session_state.side,
                st.session_state.plan, ui, st.session_state.chat_history)
            st.session_state.chat_history.append({"role": "user", "content": ui})
            st.session_state.chat_history.append({"role": "assistant", "content": result})
            st.session_state.position_results[pos] = result
        save_session()
        st.rerun()

    if st.session_state.chat_history and st.button("🗑️ 清空对话重新开始", use_container_width=True):
        st.session_state.chat_history = []
        st.rerun()


# ---- 主路由 ----
render_sidebar()
if st.session_state.page == "input":
    render_input_page()
elif st.session_state.page == "plan":
    render_plan_page()
elif st.session_state.page == "train":
    render_train_page()
